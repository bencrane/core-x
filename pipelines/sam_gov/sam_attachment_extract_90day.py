"""SAM.gov 90-day attachment — text & structured extraction engine (Stage 4, Phases 0/1/1.5/2).

Canonical implementation of `SAM_90DAY_EXTRACTION_PIPELINE_SPEC_V2.md` §4 (Phase 0), §5 (Phase 1
routing + content-canonical dedup), §6 (Phase 1.5 zip expansion), §7 (Phase 2 high-speed text pass
+ L4 spreadsheet lane + serialized .doc lane). OCR (Phase 3), embedding/IVF_PQ (Phase 4) and
structured field extraction (Phase 5) are SEPARATE artifacts (spec §17) and are NOT in this script.

ARCHITECTURE (locked — do not alter):
  * STATE is a NEW append-only event ledger `sam_attachment_extraction` (D1). The download SoR
    `sam_attachment_files` is read-only and untouched. "Current state" = the latest-terminal
    resolution view (D2). Resume = resolution view ∪ per-result JSONL checkpoint.
  * PARALLEL extract, SINGLE committing process per dataset (D3). The pipeline commits directly to R2
    with no commit_lock, so exactly one writer (the main process) ever commits; workers are pure
    compute + read-I/O and NEVER write Lance.
  * WORKER MODEL (D8/§7.1): mp.set_start_method('spawn'); daemonize (double-fork + os.setsid) BEFORE
    pool creation; ProcessPoolExecutor(initializer=_init_worker). Each worker mints its OWN
    module-global boto3 client (max_pool_connections=4) in the initializer — the sole place a client
    is born; boto3 clients are unpickleable and fork-unsafe on macOS, so they are never inherited/passed.
  * SPILLED HANDOFF (§7.2/C13): blobs stream to a SpooledTemporaryFile (RAM up to BLOB_SPILL=16MB,
    then NVMe). For PDFs > 16MB pdfium is handed the spilled FILE OBJECT directly — pdfium reads it via
    its incremental buffer reader (OS page cache), never materializing the whole byte array (#12).
  * SERIALIZED .doc LANE (§7.2/C2): LibreOffice conversion is a strict serialized path OUTSIDE the
    multiprocessing pool (soffice is single-instance-unsafe under concurrency). SOFFICE_BIN is asserted
    at startup (fail fast). A sniff pre-pass diverts rtf->striprtf, zip->python-docx, pdf->pdfium so
    only residual OLE actually hits soffice; post-convert existence/size check (exit 0 + no output =
    retriable, not a silent terminal).
  * IDEMPOTENCY (§7.6/C11): chunk writes are merge_insert on `chunk_id` (the idempotency floor). A
    result's per-result checkpoint line is written ONLY AFTER its chunks (and its ledger event) are
    durably committed to LanceDB — never before — so resume never loses chunks (#19/#20).
  * GTM SCOPE GATE (optional): Phase 1 consults `sam_attachment_gtm_scope` (built by
    sam_attachment_gtm_scope_90day.py — the "Strained Middle" mid-market cohort: NAICS set ∩ dollar band
    ∩ frequency cap). Out-of-scope resources get a terminal `skipped_out_of_scope` and never enter a lane,
    so no parse/chunk compute is spent on them. Absent table ⇒ no gate; disabled under --max-files (smoke).
    A swappable lens — re-point the cohort by rebuilding that one table, zero parser change.

Run (provision asserts soffice + creates datasets; then route -> expand -> extract):
    doppler run --project core-x --config prd -- \
      uv run --with pylance --with pyarrow --with duckdb --with boto3 --with 'psycopg[binary]' \
        --with pypdfium2 --with python-docx --with openpyxl --with xlrd --with pdfplumber \
        --with striprtf --with charset-normalizer \
      python pipelines/sam_gov/sam_attachment_extract_90day.py --phase route
    ... --phase expand
    ... --phase extract --lane L1_scope --daemon --resume
    ... --phase extract --lane L4_structured --daemon --resume
    ... --phase extract --lane L3_triage --daemon --resume

Smoke (spec §11 step 3 — 200 files to throwaway sinks; records sustained files/s, MB/s, cpu/wait):
    ... --phase extract --max-files 200 \
        --extraction-uri s3://data-sink/active/_smoke_extract_ledger/ \
        --scope-uri s3://data-sink/active/_smoke_scope/ \
        --pricing-uri s3://data-sink/active/_smoke_pricing/ \
        --unknown-uri s3://data-sink/active/_smoke_unknown/
"""

from __future__ import annotations

import argparse
import datetime as dt
import hashlib
import json
import os
import re
import sys

# ── URIs (immutable inputs read-only; new Stage-4 sinks under active/) ────────────────────────────
FILES_LEDGER_URI = os.environ.get(
    "SAM90_FILES_URI", "s3://data-sink/active/sam_attachment_files/")        # INPUT (read-only SoR)
BLOB_PREFIX = os.environ.get("SAM90_BLOB_PREFIX", "s3://data-sink/active/sam_attachment_blobs/")  # INPUT CAS
MANIFEST_URI = os.environ.get(
    "SAM90_MANIFEST_URI", "s3://data-sink/active/sam_opps_attachment_manifest_winners/")  # INPUT (award join)
EXTRACTION_URI = os.environ.get(
    "SAM90_EXTRACTION_URI", "s3://data-sink/active/sam_attachment_extraction/")  # append-only event ledger
SCOPE_URI = os.environ.get("SAM90_SCOPE_URI", "s3://data-sink/active/govcon_scope_vectors/")
PRICING_URI = os.environ.get("SAM90_PRICING_URI", "s3://data-sink/active/govcon_pricing/")
UNKNOWN_URI = os.environ.get("SAM90_UNKNOWN_URI", "s3://data-sink/active/govcon_unknown/")
DEDUP_URI = os.environ.get("SAM90_DEDUP_URI", "s3://data-sink/active/sam_attachment_content_dedup/")
INNER_URI = os.environ.get(  # Phase-1.5 materialization of expanded inner-file metadata (§6)
    "SAM90_INNER_URI", "s3://data-sink/active/sam_attachment_inner_files/")
SCOPE_GATE_URI = os.environ.get(  # GTM "Strained Middle" gate (sam_attachment_gtm_scope_90day.py). ABSENT ⇒ no gate
    "SAM90_SCOPE_GATE_URI", "s3://data-sink/active/sam_attachment_gtm_scope/")  # INPUT (read-only)
CKPT_PATH = os.environ.get("SAM90_EXTRACT_CKPT", "/tmp/sam_90day_extract_ckpt.jsonl")
LOG_PATH = os.environ.get("SAM90_EXTRACT_LOG", "/tmp/sam_90day_extract.log")
SOFFICE_BIN = os.environ.get("SOFFICE_BIN", "soffice")
FEED = "sam_attachment_extract_90day"

# ── Tunables (spec §14 defaults) ──────────────────────────────────────────────────────────────────
POOL_WORKERS = int(os.environ.get("POOL_WORKERS", "0")) or max(1, (os.cpu_count() or 4) - 2)
BIG_FILE_CONC = int(os.environ.get("BIG_FILE_CONC", "4"))          # >50MB concurrency cap (semaphore)
BIG_FILE_BYTES = 50 * 1024 * 1024
OCR_RATIO_THRESHOLD = int(os.environ.get("OCR_RATIO_THRESHOLD", "80"))
MIXED_PAGE_FRACTION = float(os.environ.get("MIXED_PAGE_FRACTION", "0.5"))
CHUNK_CHARS = int(os.environ.get("CHUNK_CHARS", "1200"))
CHUNK_OVERLAP = int(os.environ.get("CHUNK_OVERLAP", "180"))
LEDGER_FLUSH_K = int(os.environ.get("LEDGER_FLUSH_K", "500"))
CHUNK_FLUSH_M = int(os.environ.get("CHUNK_FLUSH_M", "10000"))
# Per-file extraction caps — bound worker/writer memory and keep the index clean. A handful of L4
# spreadsheets are multi-hundred-MB data dumps (one observed: 195 MB text → 128,503 chunks in a single
# worker result), which OOM-killed the bulk writer. p99 of real docs is ~1.4 MB / ~1,152 chunks, so a
# 4 MB / 4,000-chunk cap preserves >99% of documents whole and truncates only data-dump tails.
MAX_EXTRACT_CHARS = int(os.environ.get("MAX_EXTRACT_CHARS", str(4_000_000)))
MAX_CHUNKS_PER_FILE = int(os.environ.get("MAX_CHUNKS_PER_FILE", "4000"))
BLOB_SPILL = int(os.environ.get("BLOB_SPILL", str(16 * 1024 * 1024)))
ZIP_MAX_UNCOMPRESSED = int(os.environ.get("ZIP_MAX_UNCOMPRESSED", str(2 * 1024 * 1024 * 1024)))
ZIP_MAX_DEPTH = int(os.environ.get("ZIP_MAX_DEPTH", "2"))
PDFPLUMBER_MAX_PAGES = int(os.environ.get("PDFPLUMBER_MAX_PAGES", "50"))
COMPACT_TARGET_ROWS = int(os.environ.get("COMPACT_TARGET_ROWS", str(1_048_576)))

# ── Routing regex (spec §5.2 — token-boundary, NOT substring ILIKE; applied to lower(file_name)) ──
SCOPE_RX = (r"(^|[^a-z])(sow|pws|p\.?w\.?s|s\.?o\.?w|soo|statement of work|performance work statement|"
            r"scope of work|statement of objectives|specifications?|drawings?|salient charact)([^a-z]|$)")
DROP_RX = (r"(^|[^a-z])(sf ?1449|sf ?30|sf ?33|sf ?18|ppq|past performance questionnaire|"
           r"representations? and certifications?|cdrl)([^a-z]|$)")

# ── Content-truth triage regex (spec §7.4 — run on extracted text; case-insensitive: header patterns
#    are upper-cased, the labor lexicon lower-cased, both matched against the same body) ───────────
# Per-caveat control markings (spec §7.4). The ACTUAL matched caveat tokens are captured into the
# `content_marking` list<string> on each chunk/ledger row (NOT collapsed to a single flag), e.g.
# ['itar','dist_stmt_c']. Empty list = none detected in the scanned head (absence of evidence within
# the 2,000-char window, NOT a guarantee of public). Distribution Statements capture the letter (B–F).
_MARKING_PATTERNS = (
    ("cui", re.compile(r"CONTROLLED UNCLASSIFIED INFORMATION|\bCUI\b", re.IGNORECASE)),
    ("fouo", re.compile(r"FOR OFFICIAL USE ONLY|\bFOUO\b", re.IGNORECASE)),
    ("export_controlled", re.compile(r"EXPORT CONTROLLED", re.IGNORECASE)),
    ("itar", re.compile(r"\bITAR\b", re.IGNORECASE)),
    ("ear", re.compile(r"\bEAR\b", re.IGNORECASE)),
)
_DIST_STMT_RX = re.compile(r"DISTRIBUTION STATEMENT ([B-F])", re.IGNORECASE)


def _detect_markings(head: str) -> list[str]:
    """Ordered, de-duplicated control-marking caveats literally present in `head` (the document's own
    banner text). [] = none detected within the 2,000-char window (NOT proof of public)."""
    out: list[str] = []
    for name, rx in _MARKING_PATTERNS:
        if rx.search(head):
            out.append(name)
    for m in _DIST_STMT_RX.finditer(head):
        tok = f"dist_stmt_{m.group(1).lower()}"
        if tok not in out:
            out.append(tok)
    return out
SCOPE_HDR_RX = re.compile(
    r"PERFORMANCE WORK STATEMENT|STATEMENT OF WORK|STATEMENT OF OBJECTIVES|SCOPE OF WORK|\bPWS\b|"
    r"\bSOW\b|SPECIFICATIONS?|TECHNICAL REQUIREMENTS|SALIENT CHARACTERISTICS", re.IGNORECASE)
PRICING_HDR_RX = re.compile(
    r"WAGE DETERMINATION|SERVICE CONTRACT ACT|DAVIS[- ]BACON|\bSCA\b|\bWD\b|PRICE SCHEDULE|"
    r"SCHEDULE OF PRICES|\bCLIN\b", re.IGNORECASE)
BOILERPLATE_HDR_RX = re.compile(
    r"STANDARD FORM 1449|SOLICITATION/CONTRACT/ORDER FOR COMMERCIAL|AMENDMENT OF SOLICITATION|"
    r"PAST PERFORMANCE QUESTIONNAIRE|REPRESENTATIONS AND CERTIFICATIONS", re.IGNORECASE)
LABOR_LEXICON_RX = re.compile(
    r"labor categor|\bLCAT\b|\bFTE\b|headcount|clearance|certification|period of performance|"
    r"place of performance|wage|\bSCA\b|wage determination", re.IGNORECASE)

TRIAGE_HEAD_CHARS = 2000                 # §7.4: classify on the first ~2,000 normalized chars
TEXT_MIME = {"pdf", "docx", "doc", "txt"}
SHEET_MIME = {"xlsx", "xls", "xlsm", "xlsb"}
INNER_OK_MIME = {"pdf", "docx", "doc", "txt", "xlsx", "xls"}   # §6: inner files re-injected through routing

# Terminal states (spec §3.2). `routed`/`extracted_spreadsheet`/`requires_ocr` are intermediate (int).
_INTERMEDIATE = {"routed", "extracted_spreadsheet", "requires_ocr"}
# Audit-provenance states (e.g. the Phase-0 full-body marking pre-pass, sam_marking_fullbody_90day):
# they ANNOTATE a resource without superseding its extraction terminal, so they are NEVER resolution
# candidates (D2) — otherwise a newer audit event masks `extracted_*` and breaks the §12 reconcile.
_AUDIT_STATES = {"marking_fullbody"}
# Re-attemptable on resume (D2): not skipped even though some are "terminal-shaped".
_REATTEMPT = {"requires_ocr", "extract_failed", "ocr_failed"}

# Worker-process module globals (born ONLY in _init_worker under spawn — D8/C3). Never inherited.
_WORKER_S3 = None
_WORKER_BUCKET = None
_WORKER_PREFIX = None   # CAS key prefix (e.g. active/sam_attachment_blobs/) — prepended to blob_key
_BIG_SEMA = None        # cross-process cap on concurrent extraction of files > BIG_FILE_BYTES (§7.2)


# ════════════════════════════════════════════════════════════════════════ daemon / R2 / s3 helpers
def _daemonize(logpath: str) -> None:
    """Double-fork + setsid so the worker leaves the harness process group and survives a session
    resume. MUST be called BEFORE creating the pool / importing any threaded lib (boto3/lance):
    forking a multi-threaded process on macOS aborts (`+[NSNumber initialize] ... fork()`)."""
    if os.fork() > 0:
        os._exit(0)
    os.setsid()
    if os.fork() > 0:
        os._exit(0)
    f = open(logpath, "a", buffering=1)
    os.dup2(f.fileno(), 1)
    os.dup2(f.fileno(), 2)
    try:
        os.dup2(open(os.devnull).fileno(), 0)
    except OSError:
        pass


def _r2_storage_options() -> dict[str, str]:
    endpoint = os.environ.get("R2_ENDPOINT")
    account_id = os.environ.get("R2_ACCOUNT_ID")
    if not endpoint and account_id:
        endpoint = f"https://{account_id}.r2.cloudflarestorage.com"
    if not endpoint:
        raise RuntimeError("Set R2_ENDPOINT or R2_ACCOUNT_ID.")
    return {"aws_access_key_id": os.environ["R2_ACCESS_KEY_ID"],
            "aws_secret_access_key": os.environ["R2_SECRET_ACCESS_KEY"],
            "endpoint": endpoint, "region": "auto"}


def _make_s3_client(max_pool_connections: int):
    import boto3
    from botocore.config import Config
    so = _r2_storage_options()
    cfg = Config(request_checksum_calculation="when_required",
                 response_checksum_validation="when_required",
                 max_pool_connections=max_pool_connections,
                 retries={"max_attempts": 3, "mode": "standard"})
    return boto3.client("s3", endpoint_url=so["endpoint"],
                        aws_access_key_id=so["aws_access_key_id"],
                        aws_secret_access_key=so["aws_secret_access_key"],
                        region_name="auto", config=cfg)


def _split_s3(uri: str) -> tuple[str, str]:
    body = uri[len("s3://"):] if uri.startswith("s3://") else uri
    bucket, _, key = body.partition("/")
    if key and not key.endswith("/"):
        key += "/"
    return bucket, key


# ════════════════════════════════════════════════════════════════════════ single-committer lease (D3)
LEASE_PREFIX = os.environ.get("SAM90_LEASE_PREFIX", "s3://data-sink/active/_sink_leases/")
LEASE_TTL_S = int(os.environ.get("SAM90_LEASE_TTL_S", str(2 * 60 * 60)))


class SinkCommitLease:
    """Per-sink single-committer lease over an R2 conditional PUT (D3: one committing process per
    dataset). Binds every Lance COMMITTER to a sink — the extractor bulk writer, the embed writer
    (`sam_attachment_embed_90day.py`, which imports this class), and `phase_finalize` — so no two of
    them can ever commit to the same sink concurrently (the pipeline commits directly to R2 with no
    `commit_lock`, so concurrent committers race the manifest).

    Mechanics & semantics:
      * One lease object per sink: `{LEASE_PREFIX}{sink-uri-slug}.lease.json`. The slug is the full
        URI path, so smoke `_smoke_*` sink overrides lease distinct keys and never contend with prod.
      * ACQUIRE = `put_object(..., IfNoneMatch="*")` — atomic create-if-absent (R2 enforces the
        conditional write; probed live: duplicate create → HTTP 412 PreconditionFailed). Exactly one
        concurrent caller can win; losers raise RuntimeError naming the current holder. No spin or
        queueing — these writers are long batch jobs; a blocked acquire is an operator decision.
      * EXPIRY TAKEOVER: the lease body carries `expires_at` (= acquire time + ttl_s; default
        SAM90_LEASE_TTL_S = 2h). A crashed holder leaves its object behind; an acquirer that finds an
        EXPIRED (or unparseable) lease deletes it and re-runs the conditional create exactly once.
        That takeover race is still single-winner because the create stays conditional. Holders must
        size ttl_s to their worst-case wall clock (the extract phase passes 24h explicitly).
      * RELEASE deletes the object only if the body still carries this holder's random token, so a
        successor that legitimately took over after expiry is never clobbered. The GET+DELETE pair is
        not atomic; the residual window is benign — worst case a FREE lease object lingers and the
        next acquirer reclaims it via the expiry path.
      * ADVISORY: the protocol binds this repo's writers; it cannot stop an out-of-band process that
        bypasses it. Readers are never blocked. Manual unblock: delete the lease object.
    """

    def __init__(self, sink_uri: str, *, holder: str, ttl_s: int = LEASE_TTL_S) -> None:
        import uuid
        self.sink_uri = sink_uri
        slug = (sink_uri[len("s3://"):] if sink_uri.startswith("s3://") else sink_uri)
        slug = slug.strip("/").replace("/", "__")
        self.bucket, _pfx = _split_s3(LEASE_PREFIX)
        self.key = f"{_pfx}{slug}.lease.json"
        self.holder = holder
        self.ttl_s = ttl_s
        self.token = uuid.uuid4().hex
        self._s3 = None
        self._held = False

    def _client(self):
        if self._s3 is None:
            self._s3 = _make_s3_client(2)
        return self._s3

    def _try_create(self) -> bool:
        from botocore.exceptions import ClientError
        now = dt.datetime.now(dt.timezone.utc)
        body = json.dumps({
            "sink_uri": self.sink_uri, "holder": self.holder, "token": self.token,
            "acquired_at": now.isoformat(),
            "expires_at": (now + dt.timedelta(seconds=self.ttl_s)).isoformat(),
        }).encode()
        try:
            self._client().put_object(Bucket=self.bucket, Key=self.key, Body=body, IfNoneMatch="*")
            return True
        except ClientError as exc:
            status = exc.response.get("ResponseMetadata", {}).get("HTTPStatusCode")
            if status in (409, 412):       # 412 = held; 409 = concurrent conditional create in flight
                return False
            raise

    def _read(self) -> dict | None:
        from botocore.exceptions import ClientError
        try:
            obj = self._client().get_object(Bucket=self.bucket, Key=self.key)
            return json.loads(obj["Body"].read())
        except ClientError as exc:
            code = exc.response.get("Error", {}).get("Code")
            status = exc.response.get("ResponseMetadata", {}).get("HTTPStatusCode")
            if code in ("NoSuchKey", "404") or status == 404:
                return None
            raise

    def acquire(self) -> "SinkCommitLease":
        if self._try_create():
            self._held = True
            return self
        cur = self._read()
        if cur is not None:
            try:
                expired = dt.datetime.fromisoformat(cur["expires_at"]) <= dt.datetime.now(dt.timezone.utc)
            except (KeyError, TypeError, ValueError):
                expired = True                              # malformed lease == abandoned
            if expired:
                self._client().delete_object(Bucket=self.bucket, Key=self.key)
        if self._try_create():                              # single takeover retry (still conditional)
            self._held = True
            return self
        cur = self._read() or {}
        raise RuntimeError(
            f"sink commit lease for {self.sink_uri} is HELD by {cur.get('holder', '<unknown>')} "
            f"(acquired_at={cur.get('acquired_at')}, expires_at={cur.get('expires_at')}); refusing a "
            f"concurrent commit (D3). Wait for release/expiry, or delete "
            f"s3://{self.bucket}/{self.key} if the holder is known dead.")

    def release(self) -> None:
        if not self._held:
            return
        cur = self._read()
        if cur is not None and cur.get("token") == self.token:
            self._client().delete_object(Bucket=self.bucket, Key=self.key)
        self._held = False

    def __enter__(self) -> "SinkCommitLease":
        return self.acquire()

    def __exit__(self, *exc) -> None:
        self.release()


def _norm_mime(m: str | None) -> str:
    return (m or "").lstrip(".").lower()


def _sniff_mime(head: bytes) -> str | None:
    if head[:4] == b"%PDF":
        return "pdf"
    if head[:4] == b"PK\x03\x04":
        return "zip"
    if head[:4] == b"\xd0\xcf\x11\xe0":
        return "ole"
    if head[:5] == b"{\\rtf":
        return "rtf"
    if head[:3] == b"\xff\xd8\xff":
        return "jpg"
    if head[:8] == b"\x89PNG\r\n\x1a\n":
        return "png"
    if head[:6] in (b"GIF87a", b"GIF89a"):
        return "gif"
    s = head.lstrip()[:20].lower()
    if s.startswith(b"<!doctype html") or s.startswith(b"<html"):
        return "html"
    if head and b"\x00" not in head:
        return "txt"
    return None


def _dataset_exists(uri: str, so: dict) -> bool:
    import lance
    try:
        lance.dataset(uri, storage_options=so)
        return True
    except Exception:  # noqa: BLE001
        return False


# ════════════════════════════════════════════════════════════════════════ Phase 0 — dataset schemas
def _extraction_schema():
    import pyarrow as pa
    return pa.schema([
        ("resource_id", pa.string()), ("parent_resource_id", pa.string()),
        ("lane", pa.string()), ("stage", pa.string()), ("state", pa.string()),
        ("extractor", pa.string()), ("n_pages", pa.int32()), ("text_chars", pa.int64()),
        ("text_yield_ratio", pa.float64()), ("header_class", pa.string()),
        ("content_marking", pa.list_(pa.string())), ("n_chunks", pa.int32()),
        ("sha256_raw", pa.string()), ("sha256_text", pa.string()), ("codec", pa.string()),
        ("attempt", pa.int32()), ("worker_id", pa.string()), ("run_id", pa.string()),
        ("error", pa.string()),
        ("started_at", pa.timestamp("us", tz="UTC")), ("completed_at", pa.timestamp("us", tz="UTC")),
    ])


# Chunk-grain text sinks (scope/pricing/unknown): the per-chunk `text` and the pricing `cells` column
# are typed `large_string` (int64 offsets), NOT `string` (int32). At lake scale a single fragment's
# cumulative character-offset buffer exceeds the int32 ceiling (2**31 B); `compact_files()` then fails
# in default `reencode` mode with `LanceError(Arrow): Offset overflow`. int64 offsets fuse fragments
# cleanly with no `compaction_mode="try_binary_copy"` fallback. Re-materialize existing string-typed
# sinks (drop + rebuild) to adopt this — a large_string write into a string dataset is rejected by Lance.
def _scope_schema():
    import pyarrow as pa
    return pa.schema([
        ("chunk_id", pa.string()), ("resource_id", pa.string()), ("chunk_ix", pa.int32()),
        ("text", pa.large_string()), ("char_len", pa.int32()), ("header_class", pa.string()),
        ("content_marking", pa.list_(pa.string())), ("notice_id", pa.string()),
        ("solicitation_number", pa.string()), ("naics_code", pa.string()),
        ("contract_award_unique_key", pa.string()), ("source_extractor", pa.string()),
        ("reading_order_conf", pa.string()),
        ("embedding", pa.list_(pa.float32(), 1024)),   # nullable; populated in Phase 4 (NOT here)
        ("run_id", pa.string()), ("created_at", pa.timestamp("us", tz="UTC")),
    ])


def _pricing_schema():
    import pyarrow as pa
    return pa.schema([
        ("chunk_id", pa.string()), ("resource_id", pa.string()), ("chunk_ix", pa.int32()),
        ("text", pa.large_string()), ("char_len", pa.int32()), ("header_class", pa.string()),
        ("content_marking", pa.list_(pa.string())), ("notice_id", pa.string()),
        ("solicitation_number", pa.string()), ("naics_code", pa.string()),
        ("contract_award_unique_key", pa.string()), ("source_extractor", pa.string()),
        ("reading_order_conf", pa.string()), ("cells", pa.large_string()),  # cell-delimited table rows (C6); int64 offsets
        ("run_id", pa.string()), ("created_at", pa.timestamp("us", tz="UTC")),
    ])


def _unknown_schema():
    import pyarrow as pa
    return pa.schema([
        ("chunk_id", pa.string()), ("resource_id", pa.string()), ("chunk_ix", pa.int32()),
        ("text", pa.large_string()), ("char_len", pa.int32()), ("header_class", pa.string()),
        ("content_marking", pa.list_(pa.string())), ("notice_id", pa.string()),
        ("solicitation_number", pa.string()), ("naics_code", pa.string()),
        ("contract_award_unique_key", pa.string()), ("source_extractor", pa.string()),
        ("reading_order_conf", pa.string()),
        ("embedding", pa.list_(pa.float32(), 1024)),
        ("lexicon_hit", pa.bool_()),                    # C7: false rows cheaply excludable
        ("run_id", pa.string()), ("created_at", pa.timestamp("us", tz="UTC")),
    ])


def _dedup_schema():
    import pyarrow as pa
    return pa.schema([
        ("resource_id", pa.string()), ("sha256_raw", pa.string()),
        ("canonical_resource_id", pa.string()), ("is_canonical", pa.bool_()),
        ("notice_id", pa.string()), ("solicitation_number", pa.string()),
    ])


def _inner_schema():
    import pyarrow as pa
    return pa.schema([
        ("resource_id", pa.string()), ("parent_resource_id", pa.string()),
        ("inner_path", pa.string()), ("blob_key", pa.string()), ("mime_sniffed", pa.string()),
        ("content_length", pa.int64()), ("lane", pa.string()), ("sha256_raw", pa.string()),
        ("notice_id", pa.string()), ("solicitation_number", pa.string()), ("naics_code", pa.string()),
        ("run_id", pa.string()), ("created_at", pa.timestamp("us", tz="UTC")),
    ])


def _ensure_dataset(uri: str, schema, so: dict) -> bool:
    """Create an EMPTY Lance v2.1 dataset with an explicit schema if absent (idempotent). Returns
    True if newly created. Phase 0 does NOT build any index (no IVF_PQ, no chunk_id index — #21)."""
    import lance
    if _dataset_exists(uri, so):
        return False
    lance.write_dataset(schema.empty_table(), uri, mode="create",
                        data_storage_version="2.1", storage_options=so)
    print(f"  created {uri}", flush=True)
    return True


def phase0_create_datasets(so: dict, dsn: str | None = None) -> None:
    """Spec §4 / §11 step 0 (Provision): fail fast if soffice is absent, then materialize the five
    Stage-4 datasets with explicit pyarrow schemas + apply the §3.8 ops DDL. NO IVF_PQ, NO chunk_id index
    (built later, Phase 4/§11.6)."""
    _assert_soffice()                                  # §11.0 provision gate (fail fast on a soffice-less host)
    print("phase0: ensuring datasets (explicit schemas, v2.1, no indices) ...", flush=True)
    _ensure_dataset(EXTRACTION_URI, _extraction_schema(), so)
    _ensure_dataset(SCOPE_URI, _scope_schema(), so)
    _ensure_dataset(PRICING_URI, _pricing_schema(), so)
    _ensure_dataset(UNKNOWN_URI, _unknown_schema(), so)
    _ensure_dataset(DEDUP_URI, _dedup_schema(), so)
    if dsn:
        _apply_ops_ddl(dsn)


# ════════════════════════════════════════════════════════════════════════ resolution view (D2)
def _read_resolution(so: dict):
    """Return a DuckDB-resolved arrow table of the latest event per resource_id (terminal-first,
    then max attempt, then completed_at), with an `is_terminal` flag. Empty table if no ledger yet."""
    import duckdb
    import lance
    if not _dataset_exists(EXTRACTION_URI, so):
        return None
    led = lance.dataset(EXTRACTION_URI, storage_options=so).to_table(
        columns=["resource_id", "parent_resource_id", "lane", "state", "attempt", "completed_at"])
    con = duckdb.connect()
    con.register("led", led)
    inter = ",".join(f"'{s}'" for s in _INTERMEDIATE)
    audit = ",".join(f"'{s}'" for s in _AUDIT_STATES)
    return con.execute(f"""
        SELECT resource_id, parent_resource_id, lane, state, is_terminal FROM (
          SELECT resource_id, parent_resource_id, lane, state,
                 (state NOT IN ({inter})) AS is_terminal,
                 row_number() OVER (PARTITION BY resource_id
                   ORDER BY (state NOT IN ({inter})) DESC, attempt DESC, completed_at DESC) AS rn
          FROM led WHERE state NOT IN ({audit})
        ) WHERE rn = 1
    """).to_arrow_table()


def _read_scope_gate(so: dict):
    """GTM "Strained Middle" verdicts per resource_id (sam_attachment_gtm_scope, built by
    sam_attachment_gtm_scope_90day.py). None ⇒ no gate (Phase 1 routes every downloaded resource).
    When present, out-of-scope resources are diverted to a terminal `skipped_out_of_scope` BEFORE lane
    classification, so the extract worklist never spends parse/chunk compute on them. The gate is a
    swappable lens (NAICS set / dollar band / frequency cap live in the resolver, not here)."""
    import lance
    if not _dataset_exists(SCOPE_GATE_URI, so):
        return None
    return lance.dataset(SCOPE_GATE_URI, storage_options=so).to_table(
        columns=["resource_id", "gtm_scope", "scope_reason"])


# ════════════════════════════════════════════════════ id allow-list (default-OFF route/extract filter)
def _id_filter_sql(only_resource_ids: "set[str] | None", *, col: str) -> str:
    """SQL AND-fragment restricting `col` to an explicit id allow-list. None → '' (default OFF: every
    existing call path is byte-identical). Empty set → HARD RAISE (Guard #1): an empty filter would
    otherwise fall through to the full corpus — the classic "matched nothing → selected everything"
    footgun. Single quotes in ids are escaped, ids sorted for a deterministic fragment."""
    if only_resource_ids is None:
        return ""
    if not only_resource_ids:
        raise RuntimeError("id-filter resolved to an EMPTY set; refusing to run — an empty filter "
                           "would select the full corpus.")
    ids = ",".join("'" + i.replace("'", "''") + "'" for i in sorted(only_resource_ids))
    return f"AND {col} IN ({ids})"


def _assert_routed_subset(routed_ids, only_resource_ids: "set[str] | None") -> None:
    """GUARD #2: after routing, every routed resource_id MUST be in the allow-list. Raises on any leak
    — makes 'accidentally route the prime backlog into shared state' structurally impossible. No-op when
    the filter is OFF (only_resource_ids is None)."""
    if only_resource_ids is None:
        return
    leak = set(routed_ids) - only_resource_ids
    if leak:
        raise RuntimeError(f"ROUTE LEAK: {len(leak)} routed ids outside the allow-list "
                           f"(e.g. {sorted(leak)[:5]}). Aborting before any extract.")


# ════════════════════════════════════════════════════════════════════════ Phase 1 — dedup + route
def phase1_route(*, so: dict, run_id: str, max_files: int = 0, only_resource_ids: "set[str] | None" = None) -> dict:
    """Spec §5: content-canonical dedup pre-pass (§5.1) + token-boundary routing (§5.2). Writes the
    dedup fan-out map, then appends `routed`/terminal events to the append-only ledger (LEFT-ANTI-JOIN
    the resolution view → idempotent)."""
    import duckdb
    import lance
    import pyarrow as pa

    started = dt.datetime.now(dt.timezone.utc)
    files = lance.dataset(FILES_LEDGER_URI, storage_options=so).to_table(
        columns=["resource_id", "file_name", "mime_declared", "sha256", "status",
                 "notice_id", "solicitation_number", "naics_code"])
    con = duckdb.connect()
    con.register("files", files)

    # id allow-list fragments (default OFF when only_resource_ids is None; empty set raises = Guard #1).
    idf_f = _id_filter_sql(only_resource_ids, col="f.resource_id")
    idf_d = _id_filter_sql(only_resource_ids, col="d.resource_id")

    # §5.1 dedup pre-pass: canonical = min(resource_id) per raw-sha cluster.
    dedup = con.execute(f"""
        SELECT f.resource_id, f.sha256 AS sha256_raw,
               min(f.resource_id) OVER (PARTITION BY f.sha256) AS canonical_resource_id,
               f.notice_id, f.solicitation_number
        FROM files f WHERE f.status = 'downloaded' AND f.sha256 IS NOT NULL {idf_f}
    """).to_arrow_table()
    con.register("dedup", dedup)
    dmap = con.execute("""
        SELECT resource_id, sha256_raw, canonical_resource_id,
               (resource_id = canonical_resource_id) AS is_canonical, notice_id, solicitation_number
        FROM dedup
    """).to_arrow_table()
    _merge_dataset(DEDUP_URI, dmap.cast(_dedup_schema()), "resource_id", so)
    n_dup = con.execute("SELECT count(*) FROM dedup WHERE resource_id <> canonical_resource_id").fetchone()[0]
    print(f"phase1: dedup map rows={dmap.num_rows:,} non_canonical={n_dup:,}", flush=True)

    # Anti-join the resolution view (idempotent re-run).
    res = _read_resolution(so)
    if res is not None:
        con.register("res", res)
        seen = "AND f.resource_id NOT IN (SELECT resource_id FROM res)"
    else:
        seen = ""

    # GTM "Strained Middle" gate (§ de-contamination): consult the precomputed scope table. Disabled on a
    # capped smoke (max_files) to preserve raw-throughput calibration. Absent ⇒ no gate. Present ⇒ in-scope
    # (or unscored) canonical resources get a lane; out-of-scope ones are diverted to a terminal
    # `skipped_out_of_scope` below (which outranks any prior intermediate `routed` in the resolution view).
    # GATE BYPASS: an explicit id set IS the scope decision — the GTM "Strained Middle" gate must NOT
    # re-skip them (all targets are out_of_scope in that prime-cohort gate, so without this the throwaway
    # ledger would re-derive skipped_out_of_scope for the entire allow-list). Capped smoke disables it too.
    scope = None if (max_files or only_resource_ids is not None) else _read_scope_gate(so)
    if scope is not None:
        con.register("scope", scope)
        scope_join = "LEFT JOIN scope s ON f.resource_id = s.resource_id"
        scope_keep = "AND (s.resource_id IS NULL OR s.gtm_scope = 'in_scope')"
        print(f"phase1: GTM gate ON ({scope.num_rows:,} verdicts)", flush=True)
    else:
        scope_join = scope_keep = ""

    routed = con.execute(f"""
        WITH canon AS (
          SELECT f.resource_id, f.file_name, lower(coalesce(f.mime_declared,'')) AS mime,
                 f.sha256, f.notice_id, f.solicitation_number, f.naics_code
          FROM files f JOIN dedup d ON f.resource_id = d.resource_id
          {scope_join}
          WHERE f.status = 'downloaded' AND d.resource_id = d.canonical_resource_id {seen} {scope_keep} {idf_f}
        )
        SELECT resource_id, file_name, mime, sha256, notice_id, solicitation_number, naics_code,
          CASE
            WHEN mime = 'zip'                                  THEN 'container'
            WHEN mime IN ('xlsx','xls','xlsm','xlsb')          THEN 'L4_structured'
            WHEN mime NOT IN ('pdf','docx','doc','txt')        THEN 'non_text'
            WHEN regexp_matches(file_name_l, ?)                THEN 'L1_scope'
            WHEN regexp_matches(file_name_l, ?)                THEN 'L2_drop'
            ELSE 'L3_triage'
          END AS lane
        FROM (SELECT *, lower(coalesce(file_name,'')) AS file_name_l FROM canon)
    """, [SCOPE_RX, DROP_RX]).fetchall()

    # GUARD #2: before ANY extract, prove the routed parent set ⊆ the allow-list (no-op when OFF).
    _assert_routed_subset((r[0] for r in routed), only_resource_ids)

    # Non-canonical → dropped_duplicate (terminal); routed/terminal lanes per §5.2.
    noncanon = con.execute(f"""
        SELECT d.resource_id, d.notice_id, d.solicitation_number, NULL AS naics_code, d.sha256_raw
        FROM dedup d WHERE d.resource_id <> d.canonical_resource_id
        {('AND d.resource_id NOT IN (SELECT resource_id FROM res)') if res is not None else ''} {idf_d}
    """).fetchall()

    # GTM out-of-scope (terminal `skipped_out_of_scope`): downloaded canonical resources the gate rejects
    # and that are not already terminal (avoids re-skipping extracted files / duplicate skips). The reason
    # tag (out_of_scope_naics | below_band | above_band | failed_frequency_cap | no_award_link) rides `error`.
    oos = []
    if scope is not None:
        oos = con.execute(f"""
            SELECT f.resource_id, s.scope_reason, f.sha256
            FROM files f JOIN dedup d ON f.resource_id = d.resource_id
            JOIN scope s ON f.resource_id = s.resource_id
            WHERE f.status = 'downloaded' AND d.resource_id = d.canonical_resource_id
              AND s.gtm_scope <> 'in_scope'
              {('AND f.resource_id NOT IN (SELECT resource_id FROM res WHERE is_terminal)') if res is not None else ''}
        """).fetchall()

    _LANE_STATE = {"container": ("routed", "route"), "L4_structured": ("routed", "route"),
                   "non_text": ("skipped_non_text", "route"), "L1_scope": ("routed", "route"),
                   "L2_drop": ("dropped_boilerplate", "route"), "L3_triage": ("routed", "route")}
    events, counts = [], {}
    now = dt.datetime.now(dt.timezone.utc)
    for rid, _fn, _mime, sha, nid, sol, naics, lane in routed:
        state, stage = _LANE_STATE[lane]
        events.append(_ledger_row(resource_id=rid, lane=lane, stage=stage, state=state,
                                  sha256_raw=sha, run_id=run_id, started_at=now, completed_at=now))
        counts[lane] = counts.get(lane, 0) + 1
        if max_files and len(events) >= max_files:
            break
    for rid, nid, sol, _naics, sha in noncanon:
        events.append(_ledger_row(resource_id=rid, lane="dedup", stage="route",
                                  state="dropped_duplicate", sha256_raw=sha, run_id=run_id,
                                  started_at=now, completed_at=now))
    for rid, reason, sha in oos:
        events.append(_ledger_row(resource_id=rid, lane="out_of_scope", stage="route",
                                  state="skipped_out_of_scope", sha256_raw=sha, error=reason,
                                  run_id=run_id, started_at=now, completed_at=now))
    if events:
        at = pa.Table.from_pylist(events, schema=_extraction_schema())
        _append_dataset(EXTRACTION_URI, at, so)
    counts["dropped_duplicate"] = len(noncanon)
    counts["skipped_out_of_scope"] = len(oos)
    print(f"phase1: routed events={len(events):,} lanes={counts}", flush=True)
    return {"phase": "route", "lane": "all", "files_in": len(events), "lanes": counts,
            "started_at": started, "completed_at": dt.datetime.now(dt.timezone.utc)}


# ════════════════════════════════════════════════════════════════════════ Phase 1.5 — zip expansion
def phase15_expand(*, so: dict, run_id: str, max_files: int = 0,
                   only_resource_ids: "set[str] | None" = None) -> dict:
    """Spec §6/D10: stream-open each `container` (zip) from CAS in memory, content-address inner files
    by their own raw sha256, register synthetic ids `<rid>::<inner_path>` carrying parent lineage, and
    re-inject through §5.2 routing. Guards: depth ≤ ZIP_MAX_DEPTH, per-container uncompressed ceiling,
    skip encrypted entries. Parent → `expanded_container` (terminal)."""
    import io
    import zipfile

    import lance
    import pyarrow as pa

    started = dt.datetime.now(dt.timezone.utc)
    res = _read_resolution(so)
    if res is None:
        print("phase1.5: no ledger; run --phase route first.", flush=True)
        return {"phase": "expand", "lane": "container", "files_in": 0, "lanes": {},
                "started_at": started, "completed_at": dt.datetime.now(dt.timezone.utc)}
    import duckdb
    con = duckdb.connect()
    con.register("res", res)
    # Only the in-set zip containers expand (authoritative regardless of ledger; default OFF = all).
    idf_exp = _id_filter_sql(only_resource_ids, col="resource_id")
    containers = con.execute(
        f"SELECT resource_id FROM res WHERE lane='container' AND state='routed' {idf_exp}").fetchall()
    if max_files:
        containers = containers[:max_files]
    print(f"phase1.5: containers to expand={len(containers):,}", flush=True)

    s3 = _make_s3_client(8)
    bucket, key_prefix = _split_s3(BLOB_PREFIX)
    # Existing dedup shas (so inner files byte-identical to a known blob are marked non-canonical).
    known_sha: dict = {}
    if _dataset_exists(DEDUP_URI, so):
        for r in lance.dataset(DEDUP_URI, storage_options=so).to_table(
                columns=["sha256_raw", "canonical_resource_id"]).to_pylist():
            known_sha.setdefault(r["sha256_raw"], r["canonical_resource_id"])

    parent_events, inner_rows, dedup_rows, routed_events = [], [], [], []
    counts = {"expanded_container": 0, "inner_routed": 0, "inner_dup": 0, "encrypted": 0,
              "read_failed": 0, "binary": 0}
    now = dt.datetime.now(dt.timezone.utc)

    def _expand(rid: str, key: str, parent_meta: dict, depth: int) -> None:
        if depth > ZIP_MAX_DEPTH:
            return
        try:
            body = s3.get_object(Bucket=bucket, Key=key)["Body"].read()
        except Exception as exc:  # noqa: BLE001
            parent_events.append(_ledger_row(resource_id=rid, lane="container", stage="expand",
                                            state="extract_failed", error=f"fetch:{exc}", run_id=run_id,
                                            started_at=now, completed_at=now))
            return
        try:
            zf = zipfile.ZipFile(io.BytesIO(body))
        except Exception as exc:  # noqa: BLE001
            parent_events.append(_ledger_row(resource_id=rid, lane="container", stage="expand",
                                            state="extract_failed", error=f"zipopen:{exc}", run_id=run_id,
                                            started_at=now, completed_at=now))
            return
        total_unc = sum(i.file_size for i in zf.infolist())
        if total_unc > ZIP_MAX_UNCOMPRESSED:
            parent_events.append(_ledger_row(resource_id=rid, lane="container", stage="expand",
                                            state="extract_failed", error=f"zip_bomb:{total_unc}",
                                            run_id=run_id, started_at=now, completed_at=now))
            return
        for info in zf.infolist():
            if info.is_dir():
                continue
            inner_path = info.filename
            syn = f"{rid}::{inner_path}"
            try:
                raw = zf.read(info)
            except RuntimeError:               # encrypted entries raise "password required"
                counts["encrypted"] += 1
                routed_events.append(_ledger_row(resource_id=syn, parent_resource_id=rid,
                                                lane="container", stage="expand", state="extract_failed",
                                                error="zip_encrypted", run_id=run_id,
                                                started_at=now, completed_at=now))
                continue
            except Exception as exc:  # noqa: BLE001  (BadZipFile/CRC/truncated/decompress error)
                counts["read_failed"] += 1         # NOT encrypted — give it its own terminal row (§12)
                routed_events.append(_ledger_row(resource_id=syn, parent_resource_id=rid,
                                                lane="container", stage="expand", state="extract_failed",
                                                error=f"zip_read:{type(exc).__name__}", run_id=run_id,
                                                started_at=now, completed_at=now))
                continue
            sniff = _sniff_mime(raw[:512])
            sha = hashlib.sha256(raw).hexdigest()
            # Resolve sniff → LOGICAL mime (magic bytes alone can't tell docx/xlsx/pptx apart, all 'zip';
            # nor doc/xls, both 'ole'). OOXML structure disambiguates; a generic 'zip' is a true archive.
            if sniff == "zip":
                kind = _ooxml_kind(raw)              # docx | xlsx | pptx | None(generic archive)
                if kind in ("docx", "xlsx"):
                    logical = kind
                elif kind is None and depth + 1 <= ZIP_MAX_DEPTH:
                    s3.put_object(Bucket=bucket, Key=f"{key_prefix}{sha}", Body=raw)
                    _expand(syn, f"{key_prefix}{sha}", parent_meta, depth + 1)
                    continue                         # nested container → recurse
                else:
                    logical = None                   # pptx (out of v2 text scope) / too-deep archive
            elif sniff == "ole":
                # OLE magic can't distinguish doc/xls/ppt/vsd/msg — use the extension; anything not a
                # spec-admitted OLE type (doc/xls) → None → non_text (never the .doc text engine, §6).
                ext = inner_path.rsplit(".", 1)[-1].lower() if "." in inner_path else ""
                logical = {"xls": "xls", "doc": "doc"}.get(ext)
            elif sniff in ("pdf", "txt"):
                logical = sniff
            else:
                logical = None
            if logical not in INNER_OK_MIME:
                counts["binary"] += 1            # .dwg/.dgn/.rvt/images/pptx → non_text (NOT pdfium text)
                routed_events.append(_ledger_row(resource_id=syn, parent_resource_id=rid, lane="non_text",
                                                stage="expand", state="skipped_non_text", sha256_raw=sha,
                                                run_id=run_id, started_at=now, completed_at=now))
                continue
            canonical = known_sha.get(sha)
            is_canon = canonical is None
            if is_canon:
                known_sha[sha] = syn
                s3.put_object(Bucket=bucket, Key=f"{key_prefix}{sha}", Body=raw)   # content-address
            dedup_rows.append({"resource_id": syn, "sha256_raw": sha,
                               "canonical_resource_id": canonical or syn, "is_canonical": is_canon,
                               "notice_id": parent_meta.get("notice_id"),
                               "solicitation_number": parent_meta.get("solicitation_number")})
            if not is_canon:
                counts["inner_dup"] += 1
                routed_events.append(_ledger_row(resource_id=syn, parent_resource_id=rid, lane="dedup",
                                                stage="expand", state="dropped_duplicate", sha256_raw=sha,
                                                run_id=run_id, started_at=now, completed_at=now))
                continue
            lane = _route_inner(inner_path, logical)
            inner_rows.append({"resource_id": syn, "parent_resource_id": rid, "inner_path": inner_path,
                               "blob_key": sha, "mime_sniffed": logical, "content_length": len(raw),
                               "lane": lane, "sha256_raw": sha,
                               "notice_id": parent_meta.get("notice_id"),
                               "solicitation_number": parent_meta.get("solicitation_number"),
                               "naics_code": parent_meta.get("naics_code"),
                               "run_id": run_id, "created_at": now})
            state, stage = ("dropped_boilerplate", "route") if lane == "L2_drop" else ("routed", "route")
            routed_events.append(_ledger_row(resource_id=syn, parent_resource_id=rid, lane=lane,
                                            stage=stage, state=state, sha256_raw=sha, run_id=run_id,
                                            started_at=now, completed_at=now))
            counts["inner_routed"] += int(state == "routed")
        parent_events.append(_ledger_row(resource_id=rid, lane="container", stage="expand",
                                        state="expanded_container", run_id=run_id,
                                        started_at=now, completed_at=now))
        counts["expanded_container"] += 1

    # Parent lineage (notice/sol/naics) from the files ledger.
    pmeta = {r["resource_id"]: r for r in lance.dataset(FILES_LEDGER_URI, storage_options=so).to_table(
        columns=["resource_id", "notice_id", "solicitation_number", "naics_code"]).to_pylist()}
    for (rid,) in containers:
        _expand(rid, f"{key_prefix}{rid}", pmeta.get(rid, {}), depth=1)

    if dedup_rows:
        _merge_dataset(DEDUP_URI, pa.Table.from_pylist(dedup_rows, schema=_dedup_schema()),
                       "resource_id", so)
    if inner_rows:
        _merge_dataset(INNER_URI, pa.Table.from_pylist(inner_rows, schema=_inner_schema()),
                       "resource_id", so)
    allev = parent_events + routed_events
    if allev:
        _append_dataset(EXTRACTION_URI, pa.Table.from_pylist(allev, schema=_extraction_schema()), so)
    print(f"phase1.5: {counts}", flush=True)
    return {"phase": "expand", "lane": "container", "files_in": len(allev), "lanes": counts,
            "started_at": started, "completed_at": dt.datetime.now(dt.timezone.utc)}


def _ooxml_kind(raw: bytes) -> str | None:
    """Inspect a 'zip'-sniffed payload's OOXML structure: docx | xlsx | pptx | None (generic archive)."""
    import io
    import zipfile
    try:
        names = zipfile.ZipFile(io.BytesIO(raw)).namelist()
    except Exception:  # noqa: BLE001
        return None
    if any(n.startswith("xl/") for n in names):
        return "xlsx"
    if any(n.startswith("word/") for n in names):
        return "docx"
    if any(n.startswith("ppt/") for n in names):
        return "pptx"
    return None


def _route_inner(inner_path: str, mime: str) -> str:
    """Re-inject an inner file through §5.2 routing precedence (mime = resolved logical type)."""
    name = inner_path.lower()
    if mime in SHEET_MIME:
        return "L4_structured"
    if re.search(SCOPE_RX, name):
        return "L1_scope"
    if re.search(DROP_RX, name):
        return "L2_drop"
    return "L3_triage"


# ════════════════════════════════════════════════════════════════════════ ledger row factory
def _ledger_row(*, resource_id, lane, stage, state, parent_resource_id=None, extractor=None,
                n_pages=None, text_chars=None, text_yield_ratio=None, header_class=None,
                content_marking=None, n_chunks=None, sha256_raw=None, sha256_text=None, codec=None,
                attempt=1, worker_id=None, run_id=None, error=None, started_at=None,
                completed_at=None) -> dict:
    return {"resource_id": resource_id, "parent_resource_id": parent_resource_id, "lane": lane,
            "stage": stage, "state": state, "extractor": extractor, "n_pages": n_pages,
            "text_chars": text_chars, "text_yield_ratio": text_yield_ratio, "header_class": header_class,
            "content_marking": content_marking, "n_chunks": n_chunks, "sha256_raw": sha256_raw,
            "sha256_text": sha256_text, "codec": codec, "attempt": attempt, "worker_id": worker_id,
            "run_id": run_id, "error": error, "started_at": started_at, "completed_at": completed_at}


# ════════════════════════════════════════════════════════════════════════ Lance write helpers (D3)
def _append_dataset(uri: str, table, so: dict) -> None:
    import lance
    mode = "append" if _dataset_exists(uri, so) else "create"
    lance.write_dataset(table, uri, mode=mode, data_storage_version="2.1", storage_options=so)


def _merge_dataset(uri: str, table, key: str, so: dict) -> None:
    """Idempotent upsert. merge_insert if the dataset exists (created in Phase 0); else create. A
    merge_insert failure RAISES (never silently degrades to overwrite)."""
    import lance
    if _dataset_exists(uri, so):
        ds = lance.dataset(uri, storage_options=so)
        ds.merge_insert(key).when_matched_update_all().when_not_matched_insert_all().execute(table)
    else:
        lance.write_dataset(table, uri, mode="create",
                            data_storage_version="2.1", storage_options=so)


# ════════════════════════════════════════════════════════════════════════ extraction primitives
def _normalize_ws(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _chunk_text(text: str) -> list[tuple[int, str]]:
    """Spec §7.6: normalize whitespace; window CHUNK_CHARS with CHUNK_OVERLAP, snapping window ends to
    whitespace so chunks break cleanly. Deterministic (resume-safe)."""
    t = _normalize_ws(text)
    if not t:
        return []
    chunks, i, n, ix = [], 0, len(t), 0
    while i < n:
        end = min(i + CHUNK_CHARS, n)
        if end < n:
            ws = t.rfind(" ", i + max(1, CHUNK_CHARS - CHUNK_OVERLAP), end)
            if ws > i:
                end = ws
        piece = t[i:end].strip()
        if piece:
            chunks.append((ix, piece))
            ix += 1
        if end >= n or ix >= MAX_CHUNKS_PER_FILE:   # hard backstop against pathological single files
            break
        i = max(end - CHUNK_OVERLAP, i + 1)
    return chunks


def _extract_pdf(spool, content_length: int):
    """pdfium text pass. For content_length > BLOB_SPILL the spilled FILE OBJECT is handed to pdfium
    (incremental buffer reader / OS page cache) — the byte array is NOT re-read into memory (#12).
    Returns (text, n_pages, per_page_chars, reading_order_conf)."""
    import pypdfium2 as pdfium
    spool.seek(0)
    if content_length and content_length > BLOB_SPILL:
        doc = pdfium.PdfDocument(spool)            # spilled file object → page cache, no full read
    else:
        doc = pdfium.PdfDocument(spool.read())     # small → in-RAM bytes
    parts, per_page = [], []
    multi_col = False
    acc = 0
    try:
        n_pages = len(doc)
        for pi in range(n_pages):
            page = doc[pi]
            tp = page.get_textpage()
            try:
                s = tp.get_text_range() or ""
                if pi < 3 and not multi_col:
                    multi_col = _is_multicolumn(tp)
            finally:
                tp.close()
                page.close()
            parts.append(s)
            per_page.append(len(s))
            acc += len(s)
            if acc >= MAX_EXTRACT_CHARS:            # cap text; high-text PDFs never OCR, so per_page is moot
                break
    finally:
        doc.close()
    conf = "low" if multi_col else "high"          # pdfium yields stream-order, not reading-order (#15)
    return "\n".join(parts), n_pages, per_page, conf


def _is_multicolumn(textpage) -> bool:
    """Cheap x-column-cluster signal (#15): if char x-centers cluster into a left and a right band with
    a clear central gutter, the page is multi-column → reading_order_conf='low'."""
    try:
        n = textpage.count_chars()
        if n < 60:
            return False
        xs = []
        step = max(1, n // 400)
        for i in range(0, n, step):
            box = textpage.get_charbox(i)          # (left, bottom, right, top)
            xs.append((box[0] + box[2]) / 2.0)
        if not xs:
            return False
        lo, hi = min(xs), max(xs)
        if hi - lo < 1:
            return False
        mid = lo + (hi - lo) / 2.0
        band = (hi - lo) * 0.08
        gutter = sum(1 for x in xs if mid - band <= x <= mid + band)
        left = sum(1 for x in xs if x < mid - band)
        right = sum(1 for x in xs if x > mid + band)
        return gutter / len(xs) < 0.05 and left > len(xs) * 0.2 and right > len(xs) * 0.2
    except Exception:  # noqa: BLE001
        return False


def _extract_docx(spool) -> str:
    """python-docx in document order via iter_inner_content(); tables emit ' | '-cell, '\\n'-row, with
    merged-span dedup on id(cell._tc) and nested-table recursion (#6)."""
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    spool.seek(0)
    document = docx.Document(spool)
    out: list[str] = []
    acc = [0]

    def _emit_table(tbl: Table) -> None:
        for row in tbl.rows:
            seen, cells = set(), []
            for cell in row.cells:
                tc = id(cell._tc)
                if tc in seen:
                    continue
                seen.add(tc)
                cells.append(cell.text.strip())
                for blk in cell.iter_inner_content():
                    if isinstance(blk, Table):
                        _emit_table(blk)
            line = " | ".join(cells)
            out.append(line)
            acc[0] += len(line)

    for block in document.iter_inner_content():
        if acc[0] >= MAX_EXTRACT_CHARS:
            break
        if isinstance(block, Paragraph):
            out.append(block.text)
            acc[0] += len(block.text)
        elif isinstance(block, Table):
            _emit_table(block)
    return "\n".join(out)


def _extract_txt(spool):
    """Charset-aware decode (#25): utf-8 strict → BOM (utf-16/utf-8-sig) → charset-normalizer →
    cp1252(replace). Returns (text, codec, replacement_ratio)."""
    spool.seek(0)
    raw = spool.read(MAX_EXTRACT_CHARS * 4 + 8)         # bound the read (utf-8 ≤ 4 bytes/char) + BOM
    if raw[:3] == b"\xef\xbb\xbf":
        return raw.decode("utf-8-sig", "replace")[:MAX_EXTRACT_CHARS], "utf-8-sig", 0.0
    if raw[:2] in (b"\xff\xfe", b"\xfe\xff"):
        return raw.decode("utf-16", "replace")[:MAX_EXTRACT_CHARS], "utf-16", 0.0
    try:
        return raw.decode("utf-8")[:MAX_EXTRACT_CHARS], "utf-8", 0.0
    except UnicodeDecodeError:
        pass
    try:
        from charset_normalizer import from_bytes
        best = from_bytes(raw).best()
        if best is not None:
            return str(best)[:MAX_EXTRACT_CHARS], (best.encoding or "charset-normalizer"), 0.0
    except Exception:  # noqa: BLE001
        pass
    txt = raw.decode("cp1252", "replace")[:MAX_EXTRACT_CHARS]
    ratio = txt.count("�") / max(1, len(txt))
    return txt, "cp1252", ratio


def _extract_xlsx(spool) -> str:
    """openpyxl read_only — per-sheet, header + ' | '-cell, '\\n'-row delimited (NEVER a flat blob, §7.5)."""
    import openpyxl
    spool.seek(0)
    wb = openpyxl.load_workbook(spool, read_only=True, data_only=True)
    out: list[str] = []
    acc = 0
    try:
        for ws in wb.worksheets:
            if acc >= MAX_EXTRACT_CHARS:
                break
            head = f"# sheet: {ws.title}"
            out.append(head)
            acc += len(head)
            for row in ws.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(c.strip() for c in cells):
                    line = " | ".join(cells)
                    out.append(line)
                    acc += len(line)
                    if acc >= MAX_EXTRACT_CHARS:    # data-dump tail truncated; head/schema retained (§7.5)
                        break
    finally:
        wb.close()
    return "\n".join(out)


def _pdf_pricing_cells(spool) -> str | None:
    """pdfplumber table pass for pricing-classed PDF pages (C6, MIT). Bounded to PDFPLUMBER_MAX_PAGES."""
    try:
        import pdfplumber
    except Exception:  # noqa: BLE001
        return None
    spool.seek(0)
    rows: list[str] = []
    try:
        with pdfplumber.open(spool) as pdf:
            for page in pdf.pages[:PDFPLUMBER_MAX_PAGES]:
                for table in page.extract_tables() or []:
                    for r in table:
                        rows.append(" | ".join("" if c is None else str(c) for c in r))
    except Exception:  # noqa: BLE001
        return None
    return "\n".join(rows) if rows else None


# ════════════════════════════════════════════════════════════════════════ content triage (§7.4)
def _triage(text: str, lane: str) -> tuple[str, str, list[str], bool]:
    """Return (state, header_class, markings, lexicon_hit). Classify on the NORMALIZED text (§7.4)
    so header phrases broken by runs of spaces/tabs in the raw text layer still match. Control markings
    are detected FIRST (a captured list tag, not a diverting state). L1_scope bypasses the
    boilerplate-drop branch (never drop a filename-confirmed SOW)."""
    norm = _normalize_ws(text)
    head = norm[:TRIAGE_HEAD_CHARS]
    markings = _detect_markings(head)
    if SCOPE_HDR_RX.search(head):
        return "extracted_scope", "scope", markings, False
    if PRICING_HDR_RX.search(head):
        return "extracted_pricing", "pricing", markings, False
    if lane != "L1_scope" and BOILERPLATE_HDR_RX.search(head):
        return "dropped_content_noise", "boilerplate", markings, False
    # unknown: labor-lexicon admission gate over the FULL (normalized) body (#1). Hit and miss BOTH land
    # in the unknown sink (never silently dropped); miss flagged lexicon_hit=false for cheap exclusion.
    hit = bool(LABOR_LEXICON_RX.search(norm))
    return "extracted_unknown", "unknown", markings, hit


def _requires_ocr(text_chars: int, n_pages: int, per_page: list[int]) -> bool:
    """Spec §7.3 (pdf only). Low text-yield → defer to Phase 3."""
    if not n_pages:
        return text_chars == 0
    ratio = text_chars / max(1, n_pages)
    abs_floor = max(200, n_pages * 80)
    if ratio < OCR_RATIO_THRESHOLD and text_chars < abs_floor:
        return True
    sparse = sum(1 for c in per_page if c < 40)
    return (sparse / n_pages) > MIXED_PAGE_FRACTION


# ════════════════════════════════════════════════════════════════════════ worker (pure compute)
def _init_worker(big_sema=None) -> None:
    """Pool initializer (D8/C3): mint the SOLE boto3 client for this worker process as a module global
    (max_pool_connections=4). Runs under spawn, so nothing is fork-inherited. The >50MB extraction
    semaphore (§7.2) is shared in via initargs — the only safe way to share an mp primitive under spawn."""
    global _WORKER_S3, _WORKER_BUCKET, _WORKER_PREFIX, _BIG_SEMA
    _WORKER_S3 = _make_s3_client(max_pool_connections=4)
    _WORKER_BUCKET, _WORKER_PREFIX = _split_s3(BLOB_PREFIX)
    _BIG_SEMA = big_sema


def _fetch_spool(key: str):
    """Stream a CAS blob to a SpooledTemporaryFile (RAM ≤ BLOB_SPILL, then NVMe spill). `key` is the
    bare CAS address (resource_id for top-level, sha256 for expanded inner files); the blob lives under
    the BLOB_PREFIX key prefix, so it MUST be prepended to reach the object."""
    import tempfile
    body = _WORKER_S3.get_object(Bucket=_WORKER_BUCKET, Key=f"{_WORKER_PREFIX}{key}")["Body"]
    spool = tempfile.SpooledTemporaryFile(max_size=BLOB_SPILL)
    while True:
        chunk = body.read(1024 * 1024)
        if not chunk:
            break
        spool.write(chunk)
    spool.seek(0)
    return spool


def _extract_one(task: dict) -> dict:
    """Pure compute + read-I/O run in a worker. Reads the blob, dispatches by effective engine, runs
    §7.4 triage + §7.6 chunking, and RETURNS a result struct. NEVER writes Lance. Any exception is
    caught → extract_failed (the pool never dies). Files > BIG_FILE_BYTES acquire a cross-process
    semaphore so no more than BIG_FILE_CONC extract concurrently (§7.2 memory containment)."""
    import time
    started = dt.datetime.now(dt.timezone.utc)
    clen = int(task.get("content_length") or 0)
    big = clen > BIG_FILE_BYTES and _BIG_SEMA is not None
    if big:
        _BIG_SEMA.acquire()
    fetch_secs = compute_secs = 0.0
    try:
        engine = _effective_class(task["mime_declared"], task.get("mime_sniffed"),
                                  task.get("mime_match"))
        tf = time.monotonic()
        spool = _fetch_spool(task["blob_key"])
        fetch_secs = time.monotonic() - tf
        tc = time.monotonic()
        try:
            n_pages = text_chars = yield_ratio = None
            conf, codec, cells = "high", None, None
            if engine == "pdfium":
                text, n_pages, per_page, conf = _extract_pdf(spool, clen)
                text_chars = len(text)
                yield_ratio = text_chars / max(1, n_pages)
                if _requires_ocr(text_chars, n_pages, per_page):
                    res = _result_event(task, state="requires_ocr", extractor="pdfium",
                                        n_pages=n_pages, text_chars=text_chars,
                                        text_yield_ratio=yield_ratio, started=started)
                    return _timed(res, fetch_secs, time.monotonic() - tc)
            elif engine == "python_docx":
                text = _extract_docx(spool)
                text_chars = len(text)
            elif engine == "openpyxl":
                try:
                    text = _extract_xlsx(spool)
                except Exception:  # noqa: BLE001
                    # §7.5: unreadable OOXML workbook → hand to the serialized LibreOffice xlsx lane.
                    return {"fallback_xls": True, "task": task}
                text_chars = len(text)
            elif engine == "txt":
                text, codec, _ratio = _extract_txt(spool)
                text_chars = len(text)
            else:
                return _timed(_result_event(task, state="extract_failed",
                                            error=f"no_pool_engine:{engine}", started=started),
                              fetch_secs, time.monotonic() - tc)

            state, header_class, markings, lexicon_hit = _triage(text, task["lane"])
            extractor = {"pdfium": "pdfium", "python_docx": "python_docx",
                         "openpyxl": "openpyxl", "txt": "txt"}[engine]
            if state == "dropped_content_noise":
                return _timed(_result_event(task, state=state, extractor=extractor,
                                            header_class=header_class, content_marking=markings,
                                            n_pages=n_pages, text_chars=text_chars,
                                            text_yield_ratio=yield_ratio, codec=codec, started=started),
                              fetch_secs, time.monotonic() - tc)
            if state == "extracted_pricing":
                if engine == "pdfium":
                    cells = _pdf_pricing_cells(spool)
                elif engine == "openpyxl":
                    cells = text                       # the cell-delimited sheet text IS the table grid

            sink = {"extracted_scope": "scope", "extracted_pricing": "pricing",
                    "extracted_unknown": "unknown"}[state]
            sha_text = hashlib.sha256(text.encode("utf-8", "replace")).hexdigest()
            chunks = _build_chunks(task, text, header_class, markings, extractor, conf,
                                   lexicon_hit, cells, sink)
            res = _result_event(task, state=state, extractor=extractor, header_class=header_class,
                                content_marking=markings, n_pages=n_pages, text_chars=text_chars,
                                text_yield_ratio=yield_ratio, n_chunks=len(chunks),
                                sha256_text=sha_text, codec=codec, started=started,
                                sink=sink, chunks=chunks)
            return _timed(res, fetch_secs, time.monotonic() - tc)
        finally:
            spool.close()
    except Exception as exc:  # noqa: BLE001
        return _timed(_result_event(task, state="extract_failed",
                                    error=f"{type(exc).__name__}:{exc}", started=started),
                      fetch_secs, compute_secs)
    finally:
        if big:
            _BIG_SEMA.release()


def _timed(res: dict, fetch_secs: float, compute_secs: float) -> dict:
    res["fetch_secs"] = fetch_secs
    res["compute_secs"] = compute_secs
    return res


def _effective_class(mime_declared: str, mime_sniffed: str | None, mime_match) -> str:
    """Pre-dispatch sniff override (§7.2/#14): trust the sniff when the declared mime was wrong and the
    sniff is one of the known recoverable types; else dispatch by declared mime."""
    d = _norm_mime(mime_declared)
    s = _norm_mime(mime_sniffed)
    if mime_match is False and s in {"pdf", "zip", "ole", "rtf", "txt"}:
        return {"pdf": "pdfium", "zip": "python_docx", "ole": "serialized",
                "rtf": "serialized", "txt": "txt"}[s]
    if d == "pdf":
        return "pdfium"
    if d == "docx":
        return "python_docx"
    if d in SHEET_MIME:
        return "openpyxl" if d in ("xlsx", "xlsm", "xlsb") else "serialized"
    if d == "txt":
        return "txt"
    if d == "doc":
        return "serialized"
    return "serialized"


def _build_chunks(task, text, header_class, markings, extractor, conf, lexicon_hit, cells, sink):
    rid = task["resource_id"]
    now = dt.datetime.now(dt.timezone.utc)
    rows = []
    for ix, piece in _chunk_text(text):
        base = {"chunk_id": f"{rid}:{ix:04d}", "resource_id": rid, "chunk_ix": ix, "text": piece,
                "char_len": len(piece), "header_class": header_class, "content_marking": markings,
                "notice_id": task.get("notice_id"), "solicitation_number": task.get("solicitation_number"),
                "naics_code": task.get("naics_code"),
                "contract_award_unique_key": task.get("contract_award_unique_key"),
                "source_extractor": extractor, "reading_order_conf": conf,
                "run_id": task.get("run_id"), "created_at": now}
        if sink == "pricing":
            base["cells"] = cells
        else:
            base["embedding"] = None                  # nullable; populated in Phase 4
        if sink == "unknown":
            base["lexicon_hit"] = lexicon_hit
        rows.append(base)
    return rows


def _result_event(task, *, state, extractor=None, header_class=None, content_marking=None, n_pages=None,
                  text_chars=None, text_yield_ratio=None, n_chunks=None, sha256_text=None, codec=None,
                  error=None, started=None, sink=None, chunks=None) -> dict:
    completed = dt.datetime.now(dt.timezone.utc)
    led = _ledger_row(resource_id=task["resource_id"], parent_resource_id=task.get("parent_resource_id"),
                      lane=task["lane"], stage="extract", state=state, extractor=extractor,
                      n_pages=n_pages, text_chars=text_chars, text_yield_ratio=text_yield_ratio,
                      header_class=header_class, content_marking=content_marking, n_chunks=n_chunks,
                      sha256_raw=task.get("sha256_raw"), sha256_text=sha256_text, codec=codec,
                      attempt=int(task.get("attempt", 1)), worker_id=str(os.getpid()),
                      run_id=task.get("run_id"), error=error, started_at=started, completed_at=completed)
    return {"ledger": led, "sink": sink, "chunks": chunks or [], "resource_id": task["resource_id"],
            "state": state, "extractor": extractor or "none", "content_marking": content_marking,
            "text_chars": text_chars or 0, "content_length": int(task.get("content_length") or 0)}


# ════════════════════════════════════════════════════════════════════════ single-writer (D3/C11)
class _Writer:
    """The SINGLE committing process. Buffers ledger events + chunk rows; on a chunk flush it
    (1) merge_inserts chunks to each sink, (2) appends the corresponding ledger events, (3) only THEN
    writes their per-result checkpoint lines — so a checkpoint never references chunks that are not yet
    durable (§7.6/C11, #19/#20). No-chunk results (terminal drops / requires_ocr / failures) carry no
    chunk risk, so their checkpoint is written immediately."""

    def __init__(self, so: dict, ckpt_path: str):
        import pyarrow as pa
        self.so = so
        self.pa = pa
        self.ckpt = open(ckpt_path, "a", buffering=1)
        self.ledger_buf: list[dict] = []
        self.chunk_buf = {"scope": [], "pricing": [], "unknown": []}
        self.pending_ledger: list[dict] = []          # ledger events for results whose chunks are pending
        self.pending_ckpt: list[dict] = []
        self.chunk_count = 0
        self.counts = {}
        self.by_extractor = {}
        self.total_chars = 0
        self.total_chunks = 0
        self.bytes_read = 0
        self.n_results = 0
        self.cpu_secs = 0.0
        self.wait_secs = 0.0

    def _bump(self, res: dict) -> None:
        self.n_results += 1
        st = res["state"]
        self.counts[st] = self.counts.get(st, 0) + 1
        if res.get("content_marking"):                  # non-empty list => >=1 marking detected
            self.counts["content_marked"] = self.counts.get("content_marked", 0) + 1
        ext = res["extractor"]
        slot = self.by_extractor.setdefault(ext, {"ok": 0, "fail": 0})
        slot["fail" if st in ("extract_failed", "ocr_failed") else "ok"] += 1
        self.total_chars += int(res.get("text_chars") or 0)
        self.bytes_read += int(res.get("content_length") or 0)
        self.cpu_secs += float(res.get("compute_secs") or 0.0)
        self.wait_secs += float(res.get("fetch_secs") or 0.0)

    def add(self, res: dict) -> None:
        self._bump(res)
        if res["chunks"]:
            self.pending_ledger.append(res["ledger"])
            for row in res["chunks"]:
                self.chunk_buf[res["sink"]].append(row)
            self.chunk_count += len(res["chunks"])
            self.pending_ckpt.append({"resource_id": res["resource_id"], "state": res["state"],
                                      "n_chunks": len(res["chunks"])})
            if self.chunk_count >= CHUNK_FLUSH_M:
                self._flush_chunks()
        else:
            self.ledger_buf.append(res["ledger"])
            if len(self.ledger_buf) >= LEDGER_FLUSH_K:
                self._flush_ledger()
            self.ckpt.write(json.dumps({"resource_id": res["resource_id"],
                                        "state": res["state"], "n_chunks": 0}) + "\n")

    def _flush_ledger(self) -> None:
        if not self.ledger_buf:
            return
        _append_dataset(EXTRACTION_URI, self.pa.Table.from_pylist(self.ledger_buf, schema=_extraction_schema()),
                        self.so)
        self.ledger_buf.clear()

    def _flush_chunks(self) -> None:
        if not self.chunk_count:
            return
        # (1) chunks durable. APPEND (O(batch), constant) rather than per-flush merge_insert on the
        # UNINDEXED chunk_id — the latter re-scans/materializes the whole growing sink every flush
        # (O(n) → OOM at bulk scale). chunk_id uniqueness is restored by the `finalize` dedup pass;
        # deterministic chunk_id + resolution-view/checkpoint resume keep re-processing to the rare
        # crash-window, so duplicates are bounded and collapsed at the end (§12).
        for sink, rows in self.chunk_buf.items():
            if not rows:
                continue
            uri = {"scope": SCOPE_URI, "pricing": PRICING_URI, "unknown": UNKNOWN_URI}[sink]
            schema = {"scope": _scope_schema, "pricing": _pricing_schema, "unknown": _unknown_schema}[sink]()
            _append_dataset(uri, self.pa.Table.from_pylist(rows, schema=schema), self.so)
            self.total_chunks += len(rows)
            rows.clear()
        # (2) ledger events for those results durable
        if self.pending_ledger:
            _append_dataset(EXTRACTION_URI,
                            self.pa.Table.from_pylist(self.pending_ledger, schema=_extraction_schema()),
                            self.so)
            self.pending_ledger.clear()
        # (3) ONLY NOW the per-result checkpoints
        for line in self.pending_ckpt:
            self.ckpt.write(json.dumps(line) + "\n")
        self.pending_ckpt.clear()
        self.chunk_count = 0

    def finalize(self) -> None:
        self._flush_chunks()
        self._flush_ledger()
        self.ckpt.flush()
        self.ckpt.close()


# ════════════════════════════════════════════════════════════════════════ Phase 2 — text pass
def _assert_soffice() -> None:
    """Startup fail-fast (§15/C2): the serialized .doc lane requires LibreOffice."""
    import shutil
    if shutil.which(SOFFICE_BIN) is None:
        raise RuntimeError(f"SOFFICE_BIN '{SOFFICE_BIN}' not found on PATH — required for the .doc/.xls "
                           f"serialized lane (§15/C2). Set SOFFICE_BIN or install LibreOffice.")


def _build_tasks(so: dict, lanes: set[str], run_id: str,
                 only_resource_ids: "set[str] | None" = None) -> list[dict]:
    """Assemble Phase-2 work: resolution-view rows in state {routed, extract_failed} for the selected
    lane(s). Top-level files join the read-only download ledger (mime/size/lineage); inner files join
    the Phase-1.5 worklist; both carry contract_award_unique_key from the winners manifest (C14)."""
    import duckdb
    import lance
    res = _read_resolution(so)
    if res is None:
        return []
    con = duckdb.connect()
    con.register("res", res)
    lane_pred = ",".join(f"'{ln}'" for ln in lanes)
    # Belt-and-suspenders id filter (authoritative regardless of ledger). NOTE: the extract phase passes
    # None here — expanded-zip inner files carry synthetic ids `<rid>::<inner>` absent from the allow-list,
    # so the throwaway ledger (already scoped to the targets + their own inner files) IS the scope.
    extra = _id_filter_sql(only_resource_ids, col="resource_id")
    cand = con.execute(f"""
        SELECT resource_id, parent_resource_id, lane FROM res
        WHERE state IN ('routed','extract_failed') AND lane IN ({lane_pred}) {extra}
    """).to_arrow_table()
    con.register("cand", cand)

    files = lance.dataset(FILES_LEDGER_URI, storage_options=so).to_table(
        columns=["resource_id", "mime_declared", "mime_sniffed", "mime_match", "content_length",
                 "sha256", "notice_id", "solicitation_number", "naics_code"])
    con.register("files", files)
    award = lance.dataset(MANIFEST_URI, storage_options=so).to_table(
        columns=["resource_id", "contract_award_unique_key"])
    con.register("award", award)

    top = con.execute("""
        SELECT c.resource_id, c.parent_resource_id, c.lane, f.mime_declared, f.mime_sniffed,
               f.mime_match, f.content_length, f.sha256 AS sha256_raw, f.resource_id AS blob_key,
               f.notice_id, f.solicitation_number, f.naics_code, a.contract_award_unique_key
        FROM cand c JOIN files f ON c.resource_id = f.resource_id
        LEFT JOIN award a ON c.resource_id = a.resource_id
        WHERE c.parent_resource_id IS NULL
    """).to_arrow_table().to_pylist()

    inner = []
    if _dataset_exists(INNER_URI, so):
        innerds = lance.dataset(INNER_URI, storage_options=so).to_table()
        con.register("inner_wl", innerds)   # NOT "inner" — INNER is a DuckDB reserved keyword (parser error)
        inner = con.execute("""
            SELECT c.resource_id, c.parent_resource_id, c.lane,
                   i.mime_sniffed AS mime_declared, i.mime_sniffed, CAST(NULL AS BOOLEAN) AS mime_match,
                   i.content_length, i.sha256_raw, i.blob_key,
                   i.notice_id, i.solicitation_number, i.naics_code, a.contract_award_unique_key
            FROM cand c JOIN inner_wl i ON c.resource_id = i.resource_id
            LEFT JOIN award a ON c.parent_resource_id = a.resource_id
            WHERE c.parent_resource_id IS NOT NULL
        """).to_arrow_table().to_pylist()

    tasks = top + inner
    for t in tasks:
        t["run_id"] = run_id
        t["attempt"] = 1
    return tasks


def _load_checkpoint(ckpt_path: str) -> set[str]:
    """Resume done-set from the JSONL checkpoint (the resolution-view ∪ checkpoint union, §7.1). STATE-
    AWARE: a resource is "done" only if its LATEST checkpointed state is NOT re-attemptable. requires_ocr/
    extract_failed/ocr_failed (D2) must fall back to the resolution-view re-attempt — never suppressed by
    the checkpoint half of the union (last line per resource_id wins, so a later success overrides a fail)."""
    latest: dict[str, str] = {}
    if os.path.exists(ckpt_path):
        with open(ckpt_path) as fh:
            for line in fh:
                try:
                    r = json.loads(line)
                    latest[r["resource_id"]] = r.get("state", "")
                except Exception:  # noqa: BLE001
                    continue
    return {rid for rid, state in latest.items() if state not in _REATTEMPT}


def phase2_extract(*, so: dict, run_id: str, lanes: set[str], resume: bool, max_files: int = 0,
                   ckpt_path: str = CKPT_PATH, only_resource_ids: "set[str] | None" = None) -> dict:
    """Spec §7: parallel text pass (pool) + serialized .doc/.xls lane (main proc), single writer."""
    import multiprocessing as mp
    from concurrent.futures import FIRST_COMPLETED, ProcessPoolExecutor, wait

    _assert_soffice()                                  # fail fast (§15/C2)
    # D3 single-committer: hold every chunk-sink lease for the run — content-truth triage means ANY
    # lane can emit to ANY sink, so all three are bound. 24h ttl covers overnight daemonized bulk
    # passes; a crashed run's lease is reclaimed via expiry takeover (or manual delete).
    leases: list[SinkCommitLease] = []
    try:
        for _lease_uri in (SCOPE_URI, PRICING_URI, UNKNOWN_URI):
            leases.append(SinkCommitLease(_lease_uri, holder=f"extract:{run_id}",
                                          ttl_s=24 * 3600).acquire())
    except Exception:
        for lease in leases:
            lease.release()
        raise
    started = dt.datetime.now(dt.timezone.utc)
    # §B.3 caveat: the extract phase passes None to _build_tasks — the (throwaway or shared) ledger IS the
    # scope here. `only_resource_ids` is accepted for CLI symmetry but NOT forwarded: expanded-zip inner
    # files carry synthetic ids `<rid>::<inner>` absent from the allow-list, and filtering on it would drop
    # their extraction. The hard id-filter lives only at route/expand, where the files-SoR leak risk is real.
    tasks = _build_tasks(so, lanes, run_id, only_resource_ids=None)
    done = _load_checkpoint(ckpt_path) if resume else set()
    tasks = [t for t in tasks if t["resource_id"] not in done]
    if max_files:
        tasks = tasks[:max_files]

    pool_tasks, serial_tasks = [], []
    for t in tasks:
        cls = _effective_class(t["mime_declared"], t.get("mime_sniffed"), t.get("mime_match"))
        (serial_tasks if cls == "serialized" else pool_tasks).append(t)
    print(f"phase2: lanes={sorted(lanes)} pending={len(tasks):,} pool={len(pool_tasks):,} "
          f"serial(.doc/.xls)={len(serial_tasks):,} workers={POOL_WORKERS} resume_done={len(done):,}",
          flush=True)

    writer = _Writer(so, ckpt_path)
    final_status, error_text = "error", None
    import time
    t0 = time.monotonic()
    try:
        mp.set_start_method("spawn", force=True)       # D8: unpickleable/fork-unsafe boto3 clients
        ctx = mp.get_context("spawn")
        big_sema = ctx.Semaphore(BIG_FILE_CONC)        # shared in via initargs (only safe path under spawn)
        xls_fallbacks: list[dict] = []                 # OOXML workbooks openpyxl couldn't read (§7.5)
        if pool_tasks:
            with ProcessPoolExecutor(max_workers=POOL_WORKERS, initializer=_init_worker,
                                     initargs=(big_sema,), mp_context=ctx) as pool:
                it = iter(pool_tasks)
                inflight = set()
                for _ in range(POOL_WORKERS * 3):
                    try:
                        inflight.add(pool.submit(_extract_one, next(it)))
                    except StopIteration:
                        break
                while inflight:
                    fin, inflight = wait(inflight, return_when=FIRST_COMPLETED)
                    for fut in fin:
                        r = fut.result()
                        if r.get("fallback_xls"):
                            r["task"]["force_xls"] = True
                            xls_fallbacks.append(r["task"])
                        else:
                            writer.add(r)
                    for _ in range(len(fin)):
                        try:
                            inflight.add(pool.submit(_extract_one, next(it)))
                        except StopIteration:
                            break
                    if writer.n_results % 1000 == 0:
                        _progress(writer, t0)
        # Serialized .doc/.xls lane — strictly outside the pool, single instance (§7.2/C2). Also drains
        # the openpyxl→LibreOffice xlsx fallbacks (§7.5 unreadable workbooks).
        serial_all = serial_tasks + xls_fallbacks
        if serial_all:
            print(f"phase2: serialized lane — {len(serial_all):,} .doc/.xls (incl. {len(xls_fallbacks):,} "
                  f"xlsx fallbacks) via {SOFFICE_BIN}", flush=True)
            s3 = _make_s3_client(4)
            for t in serial_all:
                writer.add(_extract_serialized(t, s3, so))
                if writer.n_results % 200 == 0:
                    _progress(writer, t0)
        final_status = "success"
    except KeyboardInterrupt:
        final_status = "interrupted"
    except Exception as exc:  # noqa: BLE001
        error_text, final_status = str(exc), "error"
        print(f"FATAL: {exc}", flush=True)
    finally:
        writer.finalize()
        for lease in leases:
            try:
                lease.release()
            except Exception as exc:  # noqa: BLE001
                print(f"WARN: lease release failed for {lease.sink_uri}: {exc}", flush=True)
        elapsed = max(1e-9, time.monotonic() - t0)
        fps = writer.n_results / elapsed
        mbps = (writer.bytes_read / 1e6) / elapsed
        cpu_wait = writer.cpu_secs / max(1e-9, writer.cpu_secs + writer.wait_secs)
        print(f"phase2: results={writer.n_results:,} chunks={writer.total_chunks:,} "
              f"~{fps:.1f} files/s ~{mbps:.1f} MB/s cpu/wait={cpu_wait:.2f} status={final_status} "
              f"counts={writer.counts}", flush=True)
    return {"phase": "extract", "lane": "+".join(sorted(lanes)), "files_in": writer.n_results,
            "counts": writer.counts, "by_extractor": writer.by_extractor,
            "total_chars": writer.total_chars, "total_chunks": writer.total_chunks,
            "sustained_files_per_s": round(fps, 3), "sustained_mbps": round(mbps, 3),
            "cpu_wait_ratio": round(cpu_wait, 4), "status": final_status, "error": error_text,
            "started_at": started, "completed_at": dt.datetime.now(dt.timezone.utc)}


def _progress(writer: "_Writer", t0: float) -> None:
    import time
    el = max(1e-9, time.monotonic() - t0)
    print(f"  progress results={writer.n_results:,} chunks={writer.total_chunks:,} "
          f"~{writer.n_results/el:.1f} files/s counts={writer.counts}", flush=True)


def _extract_serialized(task: dict, s3, so: dict) -> dict:
    """SERIALIZED .doc/.xls lane (main proc, single instance). Sniff pre-pass (#8): rtf→striprtf,
    zip→python-docx, pdf→pdfium; residual OLE → soffice --convert-to pdf → pdfium (or xls → xlsx →
    openpyxl). Post-convert existence/size check: exit 0 with no output = retriable extract_failed."""
    import tempfile
    import time
    started = dt.datetime.now(dt.timezone.utc)
    tf = time.monotonic()
    fetch_secs = compute_secs = 0.0
    try:
        _bkt, _pfx = _split_s3(BLOB_PREFIX)
        body = s3.get_object(Bucket=_bkt, Key=f"{_pfx}{task['blob_key']}")["Body"].read()
        fetch_secs = time.monotonic() - tf
        tc = time.monotonic()
        sniff = _sniff_mime(body[:512]) or _norm_mime(task.get("mime_sniffed"))
        spool = tempfile.SpooledTemporaryFile(max_size=BLOB_SPILL)
        spool.write(body)
        spool.seek(0)

        if task.get("force_xls"):
            text, extractor = _soffice_xls(body, task), "libreoffice+xlsx"   # §7.5 unreadable OOXML
        elif sniff == "rtf":
            from striprtf.striprtf import rtf_to_text
            text, extractor = rtf_to_text(body.decode("latin-1", "replace")), "striprtf"
        elif sniff == "zip":
            text, extractor = _extract_docx(spool), "python_docx"
        elif sniff == "pdf":
            text, _np, _pp, _conf = _extract_pdf(spool, len(body))
            extractor = "pdfium"
        elif _looks_xls(task, sniff):
            text, extractor = _soffice_xls(body, task), "libreoffice+xlsx"
        else:
            text, extractor = _soffice_doc(body), "libreoffice+pdfium"
        spool.close()

        if text is None:
            return _timed(_result_event(task, state="extract_failed", extractor=extractor,
                                        error="soffice_no_output_retriable", started=started),
                          fetch_secs, time.monotonic() - tc)
        text_chars = len(text)
        state, header_class, markings, lexicon_hit = _triage(text, task["lane"])
        if state == "dropped_content_noise":
            return _timed(_result_event(task, state=state, extractor=extractor, header_class=header_class,
                                        content_marking=markings, text_chars=text_chars, started=started),
                          fetch_secs, time.monotonic() - tc)
        cells = text if state == "extracted_pricing" and extractor == "libreoffice+xlsx" else None
        sink = {"extracted_scope": "scope", "extracted_pricing": "pricing",
                "extracted_unknown": "unknown"}[state]
        sha_text = hashlib.sha256(text.encode("utf-8", "replace")).hexdigest()
        chunks = _build_chunks(task, text, header_class, markings, extractor, "high",
                               lexicon_hit, cells, sink)
        return _timed(_result_event(task, state=state, extractor=extractor, header_class=header_class,
                                    content_marking=markings, text_chars=text_chars, n_chunks=len(chunks),
                                    sha256_text=sha_text, started=started, sink=sink, chunks=chunks),
                      fetch_secs, time.monotonic() - tc)
    except Exception as exc:  # noqa: BLE001
        return _timed(_result_event(task, state="extract_failed",
                                    error=f"{type(exc).__name__}:{exc}", started=started),
                      fetch_secs, compute_secs)


def _looks_xls(task: dict, sniff: str | None) -> bool:
    """Legacy spreadsheet → soffice xlsx lane: declared/logical mime is a spreadsheet family member,
    or the payload sniffs OLE and the declared mime is a spreadsheet (§7.5)."""
    decl = _norm_mime(task.get("mime_declared"))
    return decl in SHEET_MIME or (sniff == "ole" and decl in SHEET_MIME)


def _soffice_doc(body: bytes) -> str | None:
    """OLE .doc → soffice --convert-to pdf → pdfium. Returns None if soffice produced no output
    (retriable). Raises on hard soffice failure."""
    import subprocess
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        src = os.path.join(td, "in.doc")
        with open(src, "wb") as fh:
            fh.write(body)
        subprocess.run([SOFFICE_BIN, "--headless", "--convert-to", "pdf", "--outdir", td, src],
                       check=True, capture_output=True, timeout=300)
        out = os.path.join(td, "in.pdf")
        if not os.path.exists(out) or os.path.getsize(out) == 0:
            return None
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(out)
        try:
            parts, acc = [], 0
            for pi in range(len(doc)):
                page = doc[pi]
                tp = page.get_textpage()
                s = tp.get_text_range() or ""
                tp.close()
                page.close()
                parts.append(s)
                acc += len(s)
                if acc >= MAX_EXTRACT_CHARS:
                    break
            return "\n".join(parts)
        finally:
            doc.close()


def _soffice_xls(body: bytes, task: dict) -> str | None:
    """Legacy .xls / unreadable spreadsheet → soffice --convert-to xlsx → openpyxl (§7.5)."""
    import subprocess
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        src = os.path.join(td, "in.xls")
        with open(src, "wb") as fh:
            fh.write(body)
        subprocess.run([SOFFICE_BIN, "--headless", "--convert-to", "xlsx", "--outdir", td, src],
                       check=True, capture_output=True, timeout=300)
        out = os.path.join(td, "in.xlsx")
        if not os.path.exists(out) or os.path.getsize(out) == 0:
            return None
        import openpyxl
        wb = openpyxl.load_workbook(out, read_only=True, data_only=True)
        try:
            lines, acc = [], 0
            for ws in wb.worksheets:
                if acc >= MAX_EXTRACT_CHARS:
                    break
                head = f"# sheet: {ws.title}"
                lines.append(head)
                acc += len(head)
                for row in ws.iter_rows(values_only=True):
                    cells = ["" if v is None else str(v) for v in row]
                    if any(c.strip() for c in cells):
                        line = " | ".join(cells)
                        lines.append(line)
                        acc += len(line)
                        if acc >= MAX_EXTRACT_CHARS:
                            break
            return "\n".join(lines)
        finally:
            wb.close()


# ════════════════════════════════════════════════════════════════════════ Phase finalize (§11.6/§12)
DEDUP_DELETE_BATCH = int(os.environ.get("DEDUP_DELETE_BATCH", "10000"))


def _duplicate_rowids(chunk_ids, rowids) -> list[int]:
    """Pure-python core of the row-address dedup: walk parallel (chunk_id, _rowid) sequences in scan
    order and return the _rowids of every DUPLICATE occurrence — the first occurrence of each
    chunk_id, in scan order, is kept. Separated from all I/O so the keep-first contract is
    unit-testable (tests/test_sam_attachment_finalize_dedup.py)."""
    seen: set = set()
    dups: list[int] = []
    for cid, rid in zip(chunk_ids, rowids):
        if cid in seen:
            dups.append(rid)
        else:
            seen.add(cid)
    return dups


def _vector_index_names(ds) -> list[str]:
    """Names of vector (ANN) indices on a dataset — [] when none. Scalar BTREE/BITMAP indices do NOT
    count: row-address deletes coexist with them; the hazard class is overwrite/compaction under an
    IVF*/HNSW index (drops/invalidates it and re-arms the rebuild cost)."""
    out: list[str] = []
    for ix in ds.list_indices():
        typ = str(ix.get("type", "")).upper()
        if "IVF" in typ or "HNSW" in typ or "VECTOR" in typ:
            out.append(str(ix.get("name", "?")))
    return out


def _assert_no_vector_index(ds, uri: str, action: str) -> None:
    """HARD GATE (build-plan anti-pattern #2): every overwrite/compaction code path against a chunk
    sink MUST call this first and MUST NOT catch the error. A `mode="overwrite"` rewrite (or
    `compact_files`) on a vector-indexed sink materializes the full table (~10 GB once embeddings
    exist) and drops every index — the silent demo-killing failure. Raises RuntimeError; never
    returns on a vector-indexed sink."""
    vec = _vector_index_names(ds)
    if vec:
        raise RuntimeError(
            f"REFUSED: {action} on {uri} — sink carries vector index(es) {vec}. An overwrite/"
            f"compaction would drop/invalidate them; use row-address delete() for dedup and the "
            f"embed module's optimize_indices() for maintenance (Phase 0 item 2 / anti-pattern #2).")


def phase_finalize(*, so: dict, run_id: str) -> dict:
    """Post-Phase-2 finalize (spec §11.6/§12) — per sink, under that sink's SinkCommitLease:

      1. ROW-ADDRESS DEDUP: collapse crash-window duplicate `chunk_id`s left by the append write path
         (restoring the uniqueness Phase-4/5 merge_insert relies on). Only the duplicate rows are
         touched: scan (chunk_id, _rowid), compute keep-first duplicates (`_duplicate_rowids`), then
         `delete("_rowid IN (...)")` in DEDUP_DELETE_BATCH-sized commits. NEVER take()+
         mode="overwrite" — that materializes the whole sink into one Arrow table (~10 GB once
         embeddings exist) and drops every index (anti-pattern #2).
      2. COMPACTION (un-indexed sinks only): compact_files + cleanup_old_versions clear the
         append-fragment debt. `_assert_no_vector_index` hard-refuses the path on any sink carrying
         a vector index — post-Phase-5 sinks get delete-based dedup ONLY here; fragment/index
         maintenance then belongs to the embed module's optimize_indices step.
      3. LEASE (D3): the per-sink SinkCommitLease guarantees finalize never commits concurrently
         with the extractor bulk writer or the embed writer; see the class docstring for semantics.

    Scalar/IVF_PQ index builds remain deferred to the post-OCR §11.6 step. Idempotent; cheap when
    there are no duplicates (a clean single-pass run)."""
    import lance
    started = dt.datetime.now(dt.timezone.utc)
    report = {}
    total_after = 0
    for name, uri in [("scope", SCOPE_URI), ("pricing", PRICING_URI), ("unknown", UNKNOWN_URI)]:
        if not _dataset_exists(uri, so):
            print(f"finalize {name}: absent, skipped", flush=True)
            continue
        with SinkCommitLease(uri, holder=f"finalize:{run_id}"):
            ds = lance.dataset(uri, storage_options=so)
            before = ds.count_rows()
            # 1. row-address dedup — scans only (chunk_id, _rowid); no full-row materialization.
            t = ds.to_table(columns=["chunk_id"], with_row_id=True)
            dup_ids = _duplicate_rowids(t.column("chunk_id").to_pylist(),
                                        t.column("_rowid").to_pylist())
            for i in range(0, len(dup_ids), DEDUP_DELETE_BATCH):
                batch = dup_ids[i:i + DEDUP_DELETE_BATCH]
                ds.delete(f"_rowid IN ({','.join(map(str, batch))})")
            if dup_ids:
                ds = lance.dataset(uri, storage_options=so)
            # 2. compaction — hard-refused on vector-indexed sinks (delete-based dedup only there).
            vec = _vector_index_names(ds)
            if vec:
                print(f"finalize {name}: vector index(es) {vec} present — compaction REFUSED; "
                      f"delete-based dedup only (run the embed module's optimize_indices instead)",
                      flush=True)
            else:
                _assert_no_vector_index(ds, uri, action="compact_files/cleanup_old_versions")
                try:
                    ds.optimize.compact_files(target_rows_per_fragment=COMPACT_TARGET_ROWS)
                    ds.cleanup_old_versions()
                except Exception as exc:  # noqa: BLE001
                    print(f"finalize {name}: compaction skipped: {exc}", flush=True)
            after = lance.dataset(uri, storage_options=so).count_rows()
        total_after += after
        report[name] = {"before": before, "dupes_removed": len(dup_ids), "after": after,
                        "vector_indexed": bool(vec)}
        print(f"finalize {name}: rows {before:,} -> {after:,} (chunk_id dupes removed "
              f"{len(dup_ids):,}); {'compaction refused (vector index)' if vec else 'compacted'}",
              flush=True)
    return {"phase": "finalize", "lane": "all", "files_in": 0, "counts": {}, "by_extractor": {},
            "total_chars": 0, "total_chunks": total_after, "sustained_files_per_s": None,
            "sustained_mbps": None, "cpu_wait_ratio": None, "status": "success", "error": None,
            "report": report, "started_at": started, "completed_at": dt.datetime.now(dt.timezone.utc)}


# ════════════════════════════════════════════════════════════════════════ ops roll-up (§3.8)
OPS_DDL = """
CREATE SCHEMA IF NOT EXISTS ops;
CREATE TABLE IF NOT EXISTS ops.sam_extraction_runs (
    id bigserial PRIMARY KEY, run_id text NOT NULL, phase text, lane text, files_in int,
    extracted_scope int, extracted_pricing int, extracted_spreadsheet int, extracted_unknown int,
    dropped_boilerplate int, dropped_duplicate int, dropped_content_noise int, content_marked int,
    expanded_container int, requires_ocr int, extract_failed int, by_extractor jsonb,
    total_chars bigint, total_chunks bigint, sustained_files_per_s numeric, sustained_mbps numeric,
    cpu_wait_ratio numeric, status text, error text, started_at timestamptz, completed_at timestamptz
);
-- Forward-rename: existing tables carried `cui_tagged`; the column is now `content_marked`
-- (count of rows with >=1 detected content marking). Idempotent.
DO $$ BEGIN
  IF EXISTS (SELECT 1 FROM information_schema.columns WHERE table_schema='ops'
             AND table_name='sam_extraction_runs' AND column_name='cui_tagged') THEN
    ALTER TABLE ops.sam_extraction_runs RENAME COLUMN cui_tagged TO content_marked;
  END IF;
END $$;
CREATE INDEX IF NOT EXISTS sam_extraction_90day_runs_run_id_idx ON ops.sam_extraction_runs (run_id);
CREATE INDEX IF NOT EXISTS sam_extraction_90day_runs_started_at_idx ON ops.sam_extraction_runs (started_at DESC);
"""


def _apply_ops_ddl(dsn: str) -> None:
    try:
        import psycopg
        with psycopg.connect(dsn) as conn, conn.cursor() as cur:
            cur.execute(OPS_DDL)
            conn.commit()
    except Exception as exc:  # noqa: BLE001
        print(f"WARN: ops DDL apply failed: {exc}", flush=True)


def _record_run(result: dict, run_id: str, dsn: str | None) -> None:
    if not dsn:
        print("WARN: no HQX_DB_URL_POOLED; skipping ops.* write.", flush=True)
        return
    c = result.get("counts", result.get("lanes", {}))
    bx = result.get("by_extractor", {})
    # extract phase: spreadsheets land as scope/pricing/unknown terminals (§7.5), so count them by their
    # extractor; route phase reports the L4 lane assignment.
    spreadsheet = (c.get("extracted_spreadsheet", 0) + c.get("L4_structured", 0)
                   + bx.get("openpyxl", {}).get("ok", 0) + bx.get("libreoffice+xlsx", {}).get("ok", 0))
    row = {
        "run_id": run_id, "phase": result["phase"], "lane": result.get("lane"),
        "files_in": result.get("files_in", 0),
        "extracted_scope": c.get("extracted_scope", 0), "extracted_pricing": c.get("extracted_pricing", 0),
        "extracted_spreadsheet": spreadsheet,
        "extracted_unknown": c.get("extracted_unknown", 0),
        "dropped_boilerplate": c.get("dropped_boilerplate", 0) + c.get("L2_drop", 0),
        "dropped_duplicate": c.get("dropped_duplicate", 0),
        "dropped_content_noise": c.get("dropped_content_noise", 0), "content_marked": c.get("content_marked", 0),
        "expanded_container": c.get("expanded_container", 0) + c.get("container", 0),
        "requires_ocr": c.get("requires_ocr", 0), "extract_failed": c.get("extract_failed", 0),
        "by_extractor": result.get("by_extractor", {}),
        "total_chars": result.get("total_chars", 0), "total_chunks": result.get("total_chunks", 0),
        "sustained_files_per_s": result.get("sustained_files_per_s"),
        "sustained_mbps": result.get("sustained_mbps"), "cpu_wait_ratio": result.get("cpu_wait_ratio"),
        "status": result.get("status", "success"), "error": result.get("error"),
        "started_at": result["started_at"], "completed_at": result["completed_at"],
    }
    try:
        import psycopg
        from psycopg.types.json import Json
        row["by_extractor"] = Json(row["by_extractor"])
        with psycopg.connect(dsn) as conn, conn.cursor() as cur:
            cur.execute(OPS_DDL)
            cur.execute("""
                INSERT INTO ops.sam_extraction_runs
                  (run_id, phase, lane, files_in, extracted_scope, extracted_pricing,
                   extracted_spreadsheet, extracted_unknown, dropped_boilerplate, dropped_duplicate,
                   dropped_content_noise, content_marked, expanded_container, requires_ocr, extract_failed,
                   by_extractor, total_chars, total_chunks, sustained_files_per_s, sustained_mbps,
                   cpu_wait_ratio, status, error, started_at, completed_at)
                VALUES (%(run_id)s,%(phase)s,%(lane)s,%(files_in)s,%(extracted_scope)s,
                   %(extracted_pricing)s,%(extracted_spreadsheet)s,%(extracted_unknown)s,
                   %(dropped_boilerplate)s,%(dropped_duplicate)s,%(dropped_content_noise)s,
                   %(content_marked)s,%(expanded_container)s,%(requires_ocr)s,%(extract_failed)s,
                   %(by_extractor)s,%(total_chars)s,%(total_chunks)s,%(sustained_files_per_s)s,
                   %(sustained_mbps)s,%(cpu_wait_ratio)s,%(status)s,%(error)s,%(started_at)s,%(completed_at)s)
                """, row)
            conn.commit()
    except Exception as exc:  # noqa: BLE001
        print(f"WARN: ops.* write failed: {exc}", flush=True)


# ════════════════════════════════════════════════════════════════════════ CLI (§17)
_PHASE_ALIASES = {"0": "phase0", "phase0": "phase0", "1": "route", "route": "route",
                  "1.5": "expand", "expand": "expand", "2": "extract", "extract": "extract",
                  "finalize": "finalize"}
_ALL_EXTRACT_LANES = ["L1_scope", "L4_structured", "L3_triage"]   # spec §11 step 4 priority order


def _cli() -> None:
    p = argparse.ArgumentParser(description="SAM.gov 90-day attachment extraction (Stage 4, Phases 0/1/1.5/2).")
    p.add_argument("--phase", required=True,
                   help="0/phase0 | 1/route | 1.5/expand | 2/extract | finalize")
    p.add_argument("--lane", default="all",
                   help="extract phase: L1_scope | L4_structured | L3_triage | all (priority order)")
    p.add_argument("--max-files", type=int, default=0, help="cap files (calibration/smoke)")
    p.add_argument("--daemon", action="store_true", help="double-fork + setsid; survive session resume")
    p.add_argument("--resume", action="store_true", help="skip resolution-view ∪ checkpoint done-set")
    p.add_argument("--run-id", default=None)
    p.add_argument("--ckpt", default=CKPT_PATH)
    # URI overrides (smoke → throwaway sinks)
    p.add_argument("--extraction-uri", default=None)
    p.add_argument("--scope-uri", default=None)
    p.add_argument("--pricing-uri", default=None)
    p.add_argument("--unknown-uri", default=None)
    p.add_argument("--dedup-uri", default=None)
    p.add_argument("--inner-uri", default=None,
                   help="Phase-1.5 inner-file worklist (expand writes it, extract reads it); "
                        "override to isolate a scoped run from the shared inner worklist")
    # id allow-list (default OFF): route/expand ONLY these resource_ids, gate forced OFF.
    p.add_argument("--resource-ids", default=None,
                   help="comma-separated ids; route/expand ONLY these (gate forced OFF)")
    p.add_argument("--resource-ids-file", default=None,
                   help="newline-delimited id file (preferred for thousands of ids)")
    a = p.parse_args()

    only_ids = None
    if a.resource_ids_file:
        with open(a.resource_ids_file) as fh:
            only_ids = {ln.strip() for ln in fh if ln.strip()}
    elif a.resource_ids:
        only_ids = {s.strip() for s in a.resource_ids.split(",") if s.strip()}

    phase = _PHASE_ALIASES.get(a.phase)
    if phase is None:
        print(f"unknown --phase {a.phase!r}; use 0|route|expand|extract", flush=True)
        sys.exit(2)
    if a.daemon:
        _daemonize(LOG_PATH)                            # BEFORE pool / threaded-lib import (D8)

    # Late URI overrides (module globals are read by every helper).
    global EXTRACTION_URI, SCOPE_URI, PRICING_URI, UNKNOWN_URI, DEDUP_URI, INNER_URI
    EXTRACTION_URI = a.extraction_uri or EXTRACTION_URI
    SCOPE_URI = a.scope_uri or SCOPE_URI
    PRICING_URI = a.pricing_uri or PRICING_URI
    UNKNOWN_URI = a.unknown_uri or UNKNOWN_URI
    DEDUP_URI = a.dedup_uri or DEDUP_URI
    INNER_URI = a.inner_uri or INNER_URI

    so = _r2_storage_options()
    dsn = os.environ.get("HQX_DB_URL_POOLED")
    run_id = a.run_id or f"90day-extract-{dt.datetime.now(dt.timezone.utc):%Y%m%dT%H%M%SZ}"

    phase0_create_datasets(so, dsn)                     # idempotent; every phase needs the sinks
    if phase == "phase0":
        print("RESULT: phase0 datasets ensured.", flush=True)
        sys.exit(0)
    if phase == "route":
        result = phase1_route(so=so, run_id=run_id, max_files=a.max_files, only_resource_ids=only_ids)
    elif phase == "expand":
        result = phase15_expand(so=so, run_id=run_id, max_files=a.max_files, only_resource_ids=only_ids)
    elif phase == "finalize":
        result = phase_finalize(so=so, run_id=run_id)
    else:
        lanes = set(_ALL_EXTRACT_LANES) if a.lane == "all" else {a.lane}
        result = phase2_extract(so=so, run_id=run_id, lanes=lanes, resume=a.resume,
                                max_files=a.max_files, ckpt_path=a.ckpt, only_resource_ids=only_ids)
    _record_run(result, run_id, dsn)
    print("RESULT:", {k: v for k, v in result.items()
                      if k not in ("started_at", "completed_at", "by_extractor")}, flush=True)
    sys.exit(0 if result.get("status", "success") in ("success",) else 1)


if __name__ == "__main__":
    _cli()
